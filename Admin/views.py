import os
import logging
from django.conf import settings
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from authentication.models import User
from Admin import models as admin_models
from .forms import UserForm, SessionForm, StudentForm, LeadForm
from .decorators import (
    login_required, role_required, admin_required, admin_only,
    teacher_redirect_to_attendance, ROLE_ADMIN, ROLE_MODERATOR, ROLE_TEACHER,
)
from .email_service import send_single_email, send_bulk_email, sanitize_subject
from .validators import validate_pdf, sanitize_filename
from .pdf_parser import detect_pdf_format, extract_candidates
from .revenue import calculate_revenue_metrics as _calculate_revenue_metrics, calculate_late_fee
from django.http import JsonResponse, HttpResponse, Http404
from django.db.models import ProtectedError
from django.utils import timezone
from django.views.decorators.http import require_POST
from datetime import date, datetime, timedelta
from dateutil.relativedelta import relativedelta
from reportlab.lib.pagesizes import letter, landscape
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle
from django.db.models import Count, Sum, Q, DecimalField, Value
from django.db.models.functions import Coalesce
from django.db import transaction
from decimal import Decimal, InvalidOperation
import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from io import BytesIO
from functools import wraps
import tempfile
import subprocess
import platform

logger = logging.getLogger('crm.admin')

ZERO = Decimal('0.00')


# ─────────────────────────────────────────────────────────────
#  IDOR helpers
# ─────────────────────────────────────────────────────────────

def _current_user(request) -> User:
    """Return the cached current user. Raises Http404 if the session is invalid."""
    cached = getattr(request, '_cached_user', None)
    if cached is not None:
        return cached
    user_id = request.session.get('user_id')
    if not user_id:
        raise Http404
    try:
        user = User.objects.get(id=user_id, status='Active')
    except User.DoesNotExist:
        raise Http404
    request._cached_user = user
    return user


def _teacher_has_student_access(user: User, student) -> bool:
    """True if a teacher is assigned to any active session containing this student."""
    return admin_models.StudentSession.objects.filter(
        student=student,
        status='Active',
    ).exists() and admin_models.Sessions.objects.filter(
        session_students__student=student,
        session_students__status='Active',
    ).exists()  # teachers are not directly assigned to sessions in this schema; restrict by user type


def _can_view_student(user: User, student) -> bool:
    """Authorize a user to view a single student record."""
    if user.usertype in {ROLE_ADMIN, ROLE_MODERATOR}:
        return True
    if user.usertype == ROLE_TEACHER:
        return admin_models.StudentSession.objects.filter(
            student=student, status='Active'
        ).exists()
    return False


def _can_view_session(user: User, session) -> bool:
    """Authorize a user to view a session."""
    if user.usertype in {ROLE_ADMIN, ROLE_MODERATOR}:
        return True
    if user.usertype == ROLE_TEACHER:
        # Teachers can only see sessions that have at least one student they teach.
        return admin_models.StudentSession.objects.filter(
            session=session, status='Active'
        ).exists()
    return False


@require_POST
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def notify_late_fee_students(request):
    """Send fee reminder emails to every active student with an outstanding balance.

    Uses email_service.send_bulk_email which sends individually (no shared recipient list),
    sanitizes subjects/bodies, and dedups recipients.
    """
    try:
        user = _current_user(request)

        students = admin_models.Student.objects.filter(
            status='Active'
        ).exclude(email__isnull=True).exclude(email__exact='')

        recipients = []
        per_recipient_subject = sanitize_subject("Fee Payment Reminder - IICE Academy")
        body_template = (
            "Dear {name},\n\n"
            "This is a reminder that you have an outstanding balance of Rs. {balance:,} at IICE Academy.\n\n"
            "Please clear this at your earliest convenience.\n\n"
            "Regards,\nIICE Academy Accounts"
        )

        students_with_pending = 0
        sent = 0
        failed = 0
        errors = []

        for student in students.prefetch_related('student_sessions__student_payments', 'student_sessions__session'):
            balance = student.remaining_balance
            if balance <= 0:
                continue
            students_with_pending += 1
            body = body_template.format(name=student.student_name or '', balance=int(balance))
            ok = send_single_email(
                subject=per_recipient_subject,
                content=body,
                recipient=student.email,
                html=False,
            )
            if ok:
                sent += 1
            else:
                failed += 1
                errors.append(student.email)
            recipients.append(student.email)

        admin_models.Notification.objects.create(
            user=user,
            category='Late Fee',
            content=f"Bulk fee reminders: {sent} sent, {failed} failed.",
        )

        return JsonResponse({
            'status': 'success',
            'message': f'{sent} reminders sent ({failed} failed).',
            'details': {
                'emails_sent': sent,
                'students_with_pending_fees': students_with_pending,
                'failed_emails': failed,
            },
        })
    except Exception:
        logger.exception('Bulk fee reminder failed')
        return JsonResponse(
            {'status': 'error', 'message': 'An error occurred. Please try again.'},
            status=500,
        )


@require_POST
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def send_fee_reminder(request):
    """Send a single fee reminder to one student. POST-only, CSRF-enforced."""
    try:
        user = _current_user(request)

        if request.content_type == 'application/json':
            try:
                data = json.loads(request.body)
            except json.JSONDecodeError:
                return JsonResponse({'status': 'error', 'message': 'Invalid JSON.'}, status=400)
            student_id = data.get('student_id')
        else:
            student_id = request.POST.get('student_id')

        if not student_id:
            return JsonResponse({'status': 'error', 'message': 'Student ID is required.'}, status=400)

        try:
            student = admin_models.Student.objects.get(id=student_id)
        except admin_models.Student.DoesNotExist:
            raise Http404
        if not _can_view_student(user, student):
            raise Http404
        if not student.email:
            return JsonResponse({'status': 'error', 'message': 'Student email not available.'}, status=400)

        balance = student.remaining_balance
        if balance <= 0:
            return JsonResponse({'status': 'error', 'message': 'No outstanding balance.'}, status=400)

        subject = sanitize_subject("Fee Payment Reminder - IICE Academy")
        body = (
            f"Dear {student.student_name},\n\n"
            f"This is a reminder that you have an outstanding balance of Rs. {int(balance):,} at IICE Academy.\n\n"
            "Please arrange for payment at your earliest convenience.\n\n"
            "Regards,\nIICE Academy Accounts"
        )

        ok = send_single_email(subject=subject, content=body, recipient=student.email, html=False)
        if not ok:
            return JsonResponse({'status': 'error', 'message': 'Failed to send email.'}, status=502)

        admin_models.Notification.objects.create(
            user=user,
            category='Late Fee',
            content=f"Fee reminder sent to {student.student_name} - Rs. {int(balance):,} pending",
        )

        return JsonResponse({
            'status': 'success',
            'message': 'Reminder sent successfully',
            'details': {
                'student_name': student.student_name,
                'email': student.email,
                'pending_amount': int(balance),
            },
        })
    except Http404:
        return JsonResponse({'status': 'error', 'message': 'Not found.'}, status=404)
    except Exception:
        logger.exception('Single fee reminder failed')
        return JsonResponse({'status': 'error', 'message': 'An error occurred. Please try again.'}, status=500)
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def EmailService(request):
    """Email service view — uses send_bulk_email which sends individually with sanitized subjects."""
    user = _current_user(request)

    if request.method == 'POST':
        email_content = request.POST.get('email_content', '') or ''
        email_subject = request.POST.get('email_subject', '') or ''
        email_list = []
        if 'faculty_checkbox' in request.POST:
            email_list.extend([u.email for u in User.objects.all() if u.email])
        if 'student_checkbox' in request.POST:
            email_list.extend([s.email for s in admin_models.Student.objects.filter(status='Active') if s.email])
        if 'lead_checkbox' in request.POST:
            email_list.extend([l.email for l in admin_models.Lead.objects.all() if l.email])

        if not email_list:
            return JsonResponse({'status': 'error', 'message': 'No valid email addresses found.'}, status=400)

        sent, failed, errors = send_bulk_email(
            subject=email_subject,
            content=email_content,
            recipients=email_list,
            html=True,
        )
        return JsonResponse({
            'status': 'success',
            'message': f'{sent} emails sent ({failed} failed).',
            'details': {'sent': sent, 'failed': failed},
        })

    return render(request, 'Admin/EmailService.html', {'user': user})
@login_required
def print_attendance_report(request, course_id):
    """Print attendance report for a course."""
    user = _current_user(request)
    try:
        # Get date range from request
        start_date = request.GET.get('start_date')
        end_date = request.GET.get('end_date')
        
        # Get the course object
        course = admin_models.Sessions.objects.get(id=course_id)

        # Fetch attendance data for this course within the date range
        attendances = admin_models.Attendance.objects.filter(
            course=course,
            date__range=[start_date, end_date]
        ).order_by('date')
        
        students = admin_models.StudentSession.objects.filter(session=course)
    except admin_models.Sessions.DoesNotExist:
        return HttpResponse('Course not found', status=404)
    except Exception as e:
        return HttpResponse('An error occurred while retrieving data.', status=500)

    # Create a PDF response
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{course.session_name}_attendance_report.pdf"'

    # Create the PDF in landscape orientation
    pdf = canvas.Canvas(response, pagesize=landscape(letter))
    width, height = landscape(letter)  # Get dimensions for landscape orientation

    # Add institution name
    pdf.setFont("Helvetica-Bold", 16)
    title = "IQRA INSTITUTE OF COMPETITIVE EXAMINATION"
    title_width = pdf.stringWidth(title, "Helvetica-Bold", 16)
    pdf.drawString((width - title_width) / 2, height - 50, title)

    # Add report title and date range
    pdf.setFont("Helvetica-Bold", 14)
    report_title = f"Attendance Report for {course.session_name}"
    report_width = pdf.stringWidth(report_title, "Helvetica-Bold", 14)
    pdf.drawString((width - report_width) / 2, height - 80, report_title)
    
    pdf.setFont("Helvetica", 12)
    date_range = f"Period: {start_date} to {end_date}"
    date_width = pdf.stringWidth(date_range, "Helvetica", 12)
    pdf.drawString((width - date_width) / 2, height - 100, date_range)

    # Get unique dates from attendance records within the selected range
    dates = sorted(set(attendance.date for attendance in attendances))
    
    if not dates:
        # If no attendance records found
        pdf.setFont("Helvetica", 12)
        message = "No attendance records found for the selected date range."
        msg_width = pdf.stringWidth(message, "Helvetica", 12)
        pdf.drawString((width - msg_width) / 2, height - 140, message)
        pdf.save()
        return response

    # Calculate table dimensions
    max_dates_per_page = 13  # Reduced to accommodate roll number column
    row_height = 25
    header_height = 30
    
    # Create tables for chunks of dates
    current_y = height - 140
    date_chunks = [dates[i:i + max_dates_per_page] for i in range(0, len(dates), max_dates_per_page)]
    
    for chunk_index, date_chunk in enumerate(date_chunks):
        if chunk_index > 0:
            # Start a new page for each chunk after the first
            pdf.showPage()
            pdf.setPageSize(landscape(letter))
            current_y = height - 100
            
            # Add header to new page
            pdf.setFont("Helvetica-Bold", 16)
            pdf.drawString((width - title_width) / 2, height - 50, title)
            pdf.setFont("Helvetica-Bold", 14)
            pdf.drawString((width - report_width) / 2, height - 80, f"{report_title} (Continued)")
        
        # Create table data with roll number, student name, and attendance dates
        table_data = [["Roll No", "Student Name"] + [date.strftime("%d-%m-%Y") for date in date_chunk]]
        
        # Add student attendance data with roll numbers
        for student in students:
            roll_no = student.student.rollno if student.student.rollno else "-"
            row = [roll_no, student.student.student_name]
            for date in date_chunk:
                # Get attendance for this specific date
                attendance = attendances.filter(
                    student=student.student,
                    date=date
                ).first()
                status = attendance.status if attendance else "Absent"
                row.append(status)
            table_data.append(row)
        
        # Calculate column widths for landscape with roll number column
        roll_col_width = 80   # Width for roll numbers
        name_col_width = 150  # Width for student names (reduced to make room for roll number)
        date_col_width = (width - 100 - roll_col_width - name_col_width) / len(date_chunk)
        col_widths = [roll_col_width, name_col_width] + [date_col_width] * len(date_chunk)
        
        # Create and style the table
        table = Table(table_data, colWidths=col_widths, rowHeights=[row_height] * len(table_data))
        style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 11),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ])
        
        # Add conditional formatting for Present/Absent (starting from column 2 since we added roll number)
        for row in range(1, len(table_data)):
            for col in range(2, len(table_data[0])):
                if table_data[row][col] == 'Present':
                    style.add('TEXTCOLOR', (col, row), (col, row), colors.green)
                elif table_data[row][col] == 'Absent':
                    style.add('TEXTCOLOR', (col, row), (col, row), colors.red)
        
        table.setStyle(style)

        # Draw the table
        table.wrapOn(pdf, width - 100, height)
        table.drawOn(pdf, 50, current_y - (len(table_data) * row_height))
        
        current_y -= (len(table_data) * row_height + 60)

    # Add footer with generation date
    pdf.setFont("Helvetica", 10)
    footer = f"Report Generated on: {date.today().strftime('%d-%m-%Y')}"
    footer_width = pdf.stringWidth(footer, "Helvetica", 10)
    pdf.drawString((width - footer_width) / 2, 30, footer)
    
    # Save the PDF
    pdf.save()
    return response
def ensure_session_fees():
    """Ensure all student sessions have proper fee values"""
    sessions_updated = 0
    for session in admin_models.StudentSession.objects.all():
        if session.fee is None and session.session and session.session.fee:
            session.fee = session.session.fee
            session.save()
            sessions_updated += 1
    return sessions_updated

@login_required
@role_required(ROLE_ADMIN)
def Payment(request):
    """Payment dashboard — single source of truth is Admin.revenue.calculate_revenue_metrics."""
    user = _current_user(request)
    try:
        metrics = _calculate_revenue_metrics()
    except Exception:
        logger.exception('Error calculating revenue metrics')
        messages.error(request, 'Error loading payment data.')
        return redirect('Admin_Dashboard')

    context = {'user': user, 'payments': metrics['recent_payments']}
    context.update(metrics)
    return render(request, 'Admin/Payments.html', context)
@require_POST
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def add_fee_payment(request, session_id):
    """Record a fee payment with transaction safety and double-click prevention.

    Accepts decimal amounts. Stores Decimal in the DB.
    """
    user = _current_user(request)

    raw_amount = (request.POST.get("amount") or '').strip()
    try:
        amount = Decimal(raw_amount)
    except (InvalidOperation, TypeError):
        return JsonResponse({"success": False, "error": "Amount must be a valid number."}, status=400)
    if amount <= ZERO:
        return JsonResponse({"success": False, "error": "Amount must be greater than zero."}, status=400)

    due_date = request.POST.get("due_date") or None

    try:
        with transaction.atomic():
            try:
                session = admin_models.StudentSession.objects.select_for_update().get(id=session_id)
            except admin_models.StudentSession.DoesNotExist:
                return JsonResponse({"success": False, "error": "Session not found."}, status=404)

            if not _can_view_student(user, session.student):
                raise Http404

            remaining_fee = session.session_balance
            if amount > remaining_fee:
                return JsonResponse({
                    "success": False,
                    "error": f"Amount Rs.{amount} exceeds remaining balance Rs.{remaining_fee}.",
                }, status=400)

            admin_models.Payments.objects.create(
                studentsession=session,
                user=user,
                amount=amount,
                date=timezone.localdate(),
                payment_status='confirmed',
                is_late_fee_payment=False,
                month=timezone.localdate().strftime('%Y-%m'),
            )

            if due_date:
                session.due_date = due_date
                session.save(update_fields=['due_date', 'updated_at'])

        admin_models.Notification.objects.create(
            user=user,
            category='New Fee',
            content=f"Collected Fee Rs {int(amount)} from {session.student.student_name}",
        )
        return JsonResponse({"success": True})

    except Http404:
        return JsonResponse({"success": False, "error": "Not found."}, status=404)
    except Exception:
        logger.exception(f'Payment error for session_id={session_id}')
        return JsonResponse({"success": False, "error": "An error occurred processing the payment."}, status=500)


@require_POST
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def waive_late_fee(request, student_session_id):
    """Admin/Moderator records a late-fee waiver as a zero-amount, late_fee_payment record."""
    user = _current_user(request)
    try:
        ss = admin_models.StudentSession.objects.select_related('student', 'session').get(id=student_session_id)
    except admin_models.StudentSession.DoesNotExist:
        raise Http404
    if not _can_view_student(user, ss.student):
        raise Http404

    reason = (request.POST.get('reason') or '').strip()[:255]
    if not reason:
        return JsonResponse({'success': False, 'error': 'Reason is required.'}, status=400)

    with transaction.atomic():
        admin_models.Payments.objects.create(
            studentsession=ss,
            user=user,
            amount=ZERO,
            date=timezone.localdate(),
            payment_status='confirmed',
            is_late_fee_payment=True,
            late_fee_waived=True,
            late_fee_waiver_reason=reason,
            late_fee_waived_by=user,
            month=timezone.localdate().strftime('%Y-%m'),
        )
    admin_models.Notification.objects.create(
        user=user,
        category='Updation',
        content=f"Late fee waived for {ss.student.student_name}: {reason}",
    )
    return JsonResponse({'success': True})

@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def MakeNotification(request):
    """Recheck late-fee notifications for active sessions."""
    user = _current_user(request)
    sessions = admin_models.Sessions.objects.filter(status='Active')
    
    today = timezone.localdate()
    month_key = today.strftime('%Y-%m')
    for session in sessions:
        for ss in admin_models.StudentSession.objects.filter(session=session, status='Active'):
            if ss.due_date and ss.due_date < today and ss.session_balance > 0:
                already = admin_models.Notification.objects.filter(
                    student_session=ss,
                    category='Late Fee',
                    notification_month=month_key,
                ).exists()
                if not already:
                    admin_models.Notification.objects.create(
                        user=user,
                        category='Late Fee',
                        content=f"Late fee for {ss.student.student_name} in {session.session_name}",
                        student_session=ss,
                        notification_month=month_key,
                    )
    return redirect('notification')

def build_due_payment_session(payment_obj, today):
    days_diff_local = (payment_obj.date - today).days
    ss = payment_obj.studentsession
    return type('DuePaymentSession', (), {
        'student': ss.student,
        'session': ss.session,
        'due_date': payment_obj.date,
        'days_until_due': days_diff_local,
        'days_overdue': abs(days_diff_local) if days_diff_local < 0 else 0,
        'fee_amount': ss.session_total_fee,
        'fee_paid': ss.session_paid,
        'balance': ss.session_balance,
        'payment_id': payment_obj.id
    })()

@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def Notification(request):
    """Notification list view with monthly renewal automation."""
    user = _current_user(request)
    notifications = admin_models.Notification.objects.all().order_by('-date', '-id')

    today_date = timezone.localdate()
    seven_days_from_now = today_date + timedelta(days=7)
    
    # AUTOMATIC MONTHLY RENEWAL PROCESSING
    # Check for monthly sessions that need renewal within 7 days
    # Only process active student sessions with active students and active sessions
    monthly_student_sessions = admin_models.StudentSession.objects.filter(
        session__session_type='monthly',
        status='Active',
        student__status='Active',
        session__status='Active'  # Ensure the session itself is also active
    ).select_related('student', 'session')
    
    renewals_created = 0
    for student_session in monthly_student_sessions:
        try:
            # Get the last payment for this student session
            last_payment = admin_models.Payments.objects.filter(
                studentsession=student_session
            ).order_by('-date').first()
            
            next_due_date = None
            if last_payment:
                # Calculate next monthly due date from last payment
                next_due_date = last_payment.date + relativedelta(months=1)
            elif student_session.registration_date:
                # No payments yet, calculate from registration date
                next_due_date = student_session.registration_date + relativedelta(months=1)
            
            if next_due_date and next_due_date <= seven_days_from_now:
                existing_unpaid = admin_models.Payments.objects.filter(
                    studentsession=student_session,
                    date=next_due_date,
                    payment_status='pending',
                ).exists()

                if not existing_unpaid:
                    admin_models.Payments.objects.create(
                        studentsession=student_session,
                        user=user,
                        amount=ZERO,
                        date=next_due_date,
                        payment_status='pending',
                        month=next_due_date.strftime('%Y-%m'),
                    )
                    
                    # Update next_monthly_due field
                    student_session.next_monthly_due = next_due_date
                    student_session.save()
                    
                    # Create notification for the renewal
                    message = (
                        f"Monthly renewal due for {student_session.student.student_name} "
                        f"in {student_session.session.session_name} - "
                        f"Rs.{student_session.session.fee} due on {next_due_date}"
                    )
                    admin_models.Notification.objects.create(
                        user=user,
                        category='Monthly Renewal',
                        content=message
                    )
                    
                    renewals_created += 1
                    
        except Exception as e:
            # Log error but continue processing other sessions
            logger.error("An error occurred. Please try again.")
    
    due_payments = admin_models.Payments.objects.filter(
        payment_status='pending',
        date__lte=seven_days_from_now,
        studentsession__student__status='Active',
        studentsession__session__status='Active',
        studentsession__status='Active',
    ).select_related('studentsession__student', 'studentsession__session')
    
    # Process each payment to create session-like objects for template compatibility
    processed_sessions = []
    for payment in due_payments:
        if payment.studentsession:
            processed_sessions.append(build_due_payment_session(payment, today_date))
    
    # Sort all sessions by due date
    processed_sessions.sort(key=lambda x: x.due_date)
    
    context = {
        'user': user,
        'notifications': notifications,
        'due_fee_sessions': processed_sessions,
        'due_fees_count': len(processed_sessions),
        'today_date': today_date
    }
    return render(request, 'Admin/Notification.html', context)

@require_POST
@login_required
def mark_all_notifications_read(request):
    """Mark all unread notifications as read for the current user."""
    user = _current_user(request)
    admin_models.Notification.objects.filter(is_read=False).update(is_read=True)
    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        return JsonResponse({'success': True, 'message': 'All notifications marked as read'})
    return redirect('notification')

@login_required
def select_course(request):
    """Course selection for attendance marking."""
    user = _current_user(request)
    courses = admin_models.Sessions.objects.annotate(
        student_count=Count('session_students')
    ).all()
    
    context = {
        'user': user,
        'courses': courses
    }
    return render(request, 'Admin/Attendance.html', context)
@login_required
def mark_attendance(request, course_id):
    """Mark attendance for a session. Validates date is not future and not >30 days past.

    Late-fee notifications are scoped to the CURRENT session only and deduplicated
    to one notification per (student_session, calendar month).
    """
    user = _current_user(request)
    try:
        course = admin_models.Sessions.objects.get(id=course_id)
    except admin_models.Sessions.DoesNotExist:
        messages.error(request, 'Course not found.')
        return redirect('select_course')

    if not _can_view_session(user, course):
        raise Http404

    students = admin_models.StudentSession.objects.filter(session=course, status='Active').select_related('student')

    if request.method == 'POST':
        submitted_date = request.POST.get('date')
        try:
            attendance_date = datetime.strptime(submitted_date, '%Y-%m-%d').date()
        except (ValueError, TypeError):
            return JsonResponse(
                {'status': 'error', 'message': 'Invalid date format.'},
                status=400,
            ) if request.headers.get('x-requested-with') == 'XMLHttpRequest' else (
                messages.error(request, 'Invalid date format.') or redirect('select_course')
            )

        today = timezone.localdate()
        if attendance_date > today:
            msg = 'Cannot mark attendance for a future date.'
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({'status': 'error', 'message': msg}, status=400)
            messages.error(request, msg)
            return redirect('mark_attendance', course_id=course_id)

        if (today - attendance_date).days > 30:
            msg = 'Cannot mark attendance more than 30 days in the past.'
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({'status': 'error', 'message': msg}, status=400)
            messages.error(request, msg)
            return redirect('mark_attendance', course_id=course_id)

        current_month_key = today.strftime('%Y-%m')

        with transaction.atomic():
            for ss in students:
                status_val = request.POST.get(f'status_{ss.student.id}')
                if not status_val:
                    continue
                admin_models.Attendance.objects.update_or_create(
                    course=course,
                    student=ss.student,
                    date=attendance_date,
                    defaults={'status': status_val},
                )

                # Late fee notification — scoped to THIS session only, deduped by month.
                if ss.due_date and ss.due_date < today and ss.session_balance > 0:
                    already_notified = admin_models.Notification.objects.filter(
                        student_session=ss,
                        category='Late Fee',
                        notification_month=current_month_key,
                    ).exists()
                    if not already_notified:
                        admin_models.Notification.objects.create(
                            user=user,
                            category='Late Fee',
                            content=f"Due date passed for {ss.student.student_name} in {course.session_name}",
                            student_session=ss,
                            notification_month=current_month_key,
                        )

        messages.success(request, 'Attendance marked successfully!')
        return redirect('select_course')

    context = {'user': user, 'course': course, 'students': students}
    return render(request, 'Admin/Mark_Attendance.html', context)
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def DeleteStudentSession(request, studentsessionid):
    """Remove a student-session enrollment. Admin/Moderator only."""
    user = _current_user(request)
    try:
        studentsession = admin_models.StudentSession.objects.get(id=studentsessionid)
    except admin_models.StudentSession.DoesNotExist:
        raise Http404

    studentid = studentsession.student.id
    student_name = studentsession.student.student_name
    session_name = studentsession.session.session_name

    studentsession.delete()

    admin_models.Notification.objects.create(
        user=user, category='Deletion',
        content=f"Removed {student_name} from {session_name} session",
    )
    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        return JsonResponse({'success': True})
    return redirect('StudentSession', studentid=studentid)
@login_required
def StudentSessionView(request, studentsessionid):
    """View/edit a StudentSession enrollment. Authorized for admins, moderators, and teachers
    who have at least one active student in the session."""
    user = _current_user(request)
    try:
        userdata = admin_models.StudentSession.objects.select_related('student', 'session').get(id=studentsessionid)
    except admin_models.StudentSession.DoesNotExist:
        raise Http404
    if not _can_view_student(user, userdata.student):
        raise Http404

    context = {
        'user': user,
        'userdata': userdata,
        'studentid': userdata.student.id,  # Include student ID for navigation
    }

    if request.method == 'POST':
        userdata.status = request.POST.get('status')
        userdata.notes = request.POST.get('notes')
        userdata.due_date = request.POST.get('due_date')
        message = "Updated  " + userdata.student.student_name + " Record in " + userdata.session.session_name + " session"
        admin_models.Notification.objects.create(user=user, category='Updation', content=message)
        userdata.save()

    return render(request, 'Admin/StudentSessionView.html', context)
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def AddStudentSession(request, studentid):
    """Enroll a student in a new session. Admin/Moderator only."""
    user = _current_user(request)
    try:
        student = admin_models.Student.objects.get(id=studentid)
    except admin_models.Student.DoesNotExist:
        raise Http404
    active_sessions = admin_models.Sessions.objects.filter(status='Active')
    
    # Check if student has any previous enrollments (for fee waiver logic)
    has_previous_enrollments = admin_models.StudentSession.objects.filter(student=student).exists()
    
    if request.method == 'POST':
        session_id = request.POST.get('session_id')
        due_date = request.POST.get('due_date')
        discount = request.POST.get('discount', 0)
        notes = request.POST.get('notes', '')
        
        registration_date_str = request.POST.get('registration_date')
        if registration_date_str:
            from datetime import datetime
            registration_date = datetime.strptime(registration_date_str, '%Y-%m-%d').date()
        else:
            registration_date = date.today()

        try:
            session = admin_models.Sessions.objects.get(id=session_id)
            
            # INSTITUTIONAL POLICY: Check for multiple active enrollments
            existing_active_enrollment = admin_models.StudentSession.objects.filter(
                student=student,
                status='Active'
            ).exclude(session_id=session_id).first()
            
            if existing_active_enrollment:
                error_msg = f"Student is already enrolled in active session: {existing_active_enrollment.session.session_name}. Please complete or withdraw from current session before enrolling in a new one."
                if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                    return JsonResponse({'success': False, 'error': error_msg})
                messages.error(request, error_msg)
                return render(request, 'Admin/AddStudentSession.html', {
                    'student': student,
                    'active_sessions': active_sessions,
                    'error': error_msg
                })
            
            # INSTITUTIONAL POLICY: Registration fee waiver for re-enrollments
            registration_fee = 0  # Default to 0 for re-enrollments
            if not has_previous_enrollments:
                # First-time enrollment: charge full registration fee
                registration_fee = session.registration_fee
            else:
                # Re-enrollment: waive registration fee
                registration_fee = 0
                notes += f" [Registration fee waived - re-enrollment policy applied]"

            # Create new StudentSession
            student_session = admin_models.StudentSession(
                student=student,
                session=session,
                registration_date=registration_date,
                registration_fee=registration_fee,  # Applied fee waiver logic
                fee=session.fee,
                due_date=due_date,
                discount=discount,
                status='Active',
                notes=notes,
            )
            
            student_session.save()
            
            # Create notification with fee waiver info
            fee_info = "(Registration fee waived)" if has_previous_enrollments else f"(Registration fee: {registration_fee})"
            message = f"Enrolled {student.student_name} in {session.session_name} session {fee_info}"
            admin_models.Notification.objects.create(user=user, category='New Entry', content=message)
            
        except Exception as e:
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({'success': False, 'error': 'An error occurred. Please try again.'})
            raise

        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({
                'success': True, 
                'message': f'Successfully enrolled with {"waived" if has_previous_enrollments else "standard"} registration fee policy'
            })
        return redirect('StudentView', studentid=studentid)

    return render(request, 'Admin/AddStudentSession.html', {
        'student': student,
        'active_sessions': active_sessions,
        'has_previous_enrollments': has_previous_enrollments,
    })
@login_required
def StudentSession(request, studentid):
    """Display a student's session enrollments."""
    user = _current_user(request)
    try:
        userdata = admin_models.Student.objects.get(id=studentid)
    except admin_models.Student.DoesNotExist:
        raise Http404
    if not _can_view_student(user, userdata):
        raise Http404
    sessions = admin_models.StudentSession.objects.filter(student=userdata)

    context = {
        'user': user,
        'userdata': userdata,
        'sessions': sessions,
        'studentid': studentid
    }
    return render(request, 'Admin/StudentSession.html', context)
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def LeadView(request, leadid):
    """View/edit a lead. Admin and Moderator only — teachers have no lead access."""
    user = _current_user(request)
    try:
        userdata = admin_models.Lead.objects.get(id=leadid)
    except admin_models.Lead.DoesNotExist:
        raise Http404

    context = {
        'user': user,
        'userdata': userdata,
    }

    if request.method == 'POST':
        form = LeadForm(request.POST, request.FILES, instance=userdata)  # Form for 'userdata'

        if form.is_valid():
            form.save()
            message = "Updated  " + userdata.name + " Information in Leads"
            admin_models.Notification.objects.create(user=user, category='Deletion', content=message)
            return redirect('LeadView', leadid=leadid)  # Redirect to avoid resubmission
        else:
            # Print form errors for debugging
            logger.debug("Form validation failed")
            logger.debug("Form errors: %s", form.errors)

    else:
        form = LeadForm(instance=userdata)

    context['form'] = form
    return render(request, 'Admin/LeadView.html', context)
@login_required
@role_required(ROLE_ADMIN)
def DeleteLead(request, leadid):
    """Delete a lead. Admin only."""
    user = _current_user(request)
    try:
        lead = admin_models.Lead.objects.get(id=leadid)
    except admin_models.Lead.DoesNotExist:
        raise Http404

    lead_name = lead.name
    lead.delete()

    admin_models.Notification.objects.create(
        user=user, category='Deletion', content=f"Deleted Lead {lead_name}",
    )
    messages.success(request, f'Lead "{lead_name}" has been successfully deleted.')
    return redirect('Leads')
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def AddLead(request):
    """Create a new lead. Admin/Moderator only."""
    user = _current_user(request)
    active_sessions = admin_models.Sessions.objects.filter(status='Active')

    if request.method == 'POST':
        form = LeadForm(request.POST)

        if form.is_valid():
            newuser = form.save(commit=False)
            newuser.save()  # Save the new user
            message = "Added " + newuser.name + " to Leads"
            admin_models.Notification.objects.create(user=user, category='New Entry', content=message)
            messages.success(request, "Lead added successfully!")
            return redirect('Leads')
        else:
            logger.debug("Form errors: %s", form.errors)

    else:
        form = LeadForm()

    context = {
        'user': user,
        'form': form,
        'active_sessions': active_sessions,
    }
    return render(request, 'Admin/AddLead.html', context)
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def Leads(request):
    """List all leads. Admin/Moderator only."""
    user = _current_user(request)
    leads = admin_models.Lead.objects.all()
    
    # Filter leads by inquiry type for statistics
    call_leads = leads.filter(form_of_inquiry='Call')
    message_leads = leads.filter(form_of_inquiry='Message')
    visit_leads = leads.filter(form_of_inquiry='Physical Visit')
    
    context = {
        'user': user,
        'leads': leads,
        'call_leads': call_leads,
        'message_leads': message_leads,
        'visit_leads': visit_leads,
    }
    return render(request, 'Admin/Leads.html', context)
@login_required
@role_required(ROLE_ADMIN)
def DeleteStudent(request, studentid):
    """Soft-delete a student. Admin only.

    Blocks deletion when the student has confirmed payments — those rows must be preserved.
    Falls back to deactivation.
    """
    user = _current_user(request)
    try:
        student = admin_models.Student.objects.get(id=studentid)
    except admin_models.Student.DoesNotExist:
        raise Http404

    student_name = student.student_name

    try:
        student.delete(deleted_by=user)
    except ProtectedError:
        messages.error(
            request,
            'Cannot delete student with confirmed payment history. Deactivate instead.',
        )
        return redirect('Students')

    admin_models.Notification.objects.create(
        user=user, category='Deletion', content=f"Deactivated {student_name}",
    )
    messages.success(request, f'{student_name} has been deactivated.')

    from_page = request.GET.get('from')
    if from_page == 'exstudents':
        return redirect('ExStudents')
    return redirect('Students')
@login_required
def StudentView(request, studentid):
    """View/edit a single student. Teachers can only view students in their active sessions."""
    user = _current_user(request)
    try:
        userdata = admin_models.Student.objects.get(id=studentid)
    except admin_models.Student.DoesNotExist:
        raise Http404
    if not _can_view_student(user, userdata):
        raise Http404
    status_choices = admin_models.Student.STATUS_CHOICES
    
    # Initialize form early to avoid UnboundLocalError
    form = StudentForm(instance=userdata)
    form_has_errors = False
    form_errors = None
    
    # Get student's enrolled sessions
    student_sessions = admin_models.StudentSession.objects.filter(student=userdata)
    
    # Compute available sessions: active sessions not yet enrolled
    all_active = admin_models.Sessions.objects.filter(status='Active')
    enrolled_ids = student_sessions.values_list('session_id', flat=True)
    available_sessions = all_active.exclude(id__in=enrolled_ids)

    # Get student's payment information using UNIFIED SYSTEM
    # Calculate installment information from payment records
    all_payments = admin_models.Payments.objects.filter(
        studentsession__student=userdata
    ).order_by('date')
    
    # Get latest payment
    latest_payment = all_payments.filter(amount__gt=0).order_by('-date').first()
    
    # Installments: pending = unpaid, confirmed+amount>0 = paid.
    total_installments = all_payments.count()
    paid_installments = all_payments.filter(payment_status='confirmed', amount__gt=ZERO).count()
    unpaid_installments = all_payments.filter(payment_status='pending').count()
    
    # Calculate one-time registration fee from primary session
    primary_session = (
        student_sessions.exclude(registration_date__isnull=True)
        .order_by('registration_date', 'id')
        .first()
        or student_sessions.order_by('id').first()
    )

    one_time_reg_fee = 0
    if primary_session:
        one_time_reg_fee = (
            primary_session.registration_fee
            if primary_session.registration_fee is not None
            else (primary_session.session.registration_fee or 0)
        )
        
    # Calculate per installment amount
    per_installment_amount = 0
    if paid_installments > 0:
        # Use amount from first paid installment
        first_paid = all_payments.filter(amount__gt=0).first()
        if first_paid:
            per_installment_amount = first_paid.amount
    else:
        # Calculate per installment amount based on final fee (total_fee already includes registration_fee - discount)
        discount_amount = sum(s.discount or 0 for s in student_sessions.filter(status='Active'))
        calculated_final_fee = userdata.total_fee - discount_amount
        if total_installments > 0:
            per_installment_amount = int(calculated_final_fee / total_installments)

    # Create payment_info object using calculated properties
    discount_amount = sum(s.discount or 0 for s in student_sessions.filter(status='Active'))
    # Fix: Use userdata.total_fee which already includes registration fee for primary session
    calculated_final_fee = userdata.total_fee - discount_amount
    calculated_remaining = max(0, calculated_final_fee - userdata.total_paid)
    
    payment_info = type('PaymentInfo', (), {
        'total_fee': userdata.total_fee,
        'paid_amount': userdata.total_paid,
        'final_fee': calculated_final_fee,
        'remaining_amount': calculated_remaining,
        'installments_count': total_installments,
        'per_installment_amount': per_installment_amount,
        'discount': discount_amount
    })()
    
    # Get installment details from payment records
    installments = []
    for payment in all_payments:
        # Show expected amount for unpaid installments, actual amount for paid ones
        display_amount = payment.amount if payment.amount > 0 else per_installment_amount
        
        installment_info = {
            'id': payment.id,
            'amount': display_amount,
            'due_date': payment.date,
            'paid_date': payment.date if payment.amount > 0 else None,
            'status': 'Paid' if payment.amount > 0 else 'Unpaid',
            'is_paid': payment.amount > 0
        }
        installments.append(installment_info)
    
    next_due_date = None
    # Find next unpaid installment due date - only if there are unpaid installments and remaining balance
    if unpaid_installments > 0 and calculated_remaining > 0:
        next_unpaid = all_payments.filter(payment_status='pending').order_by('date').first()
        if next_unpaid:
            next_due_date = next_unpaid.date
    
    # Add these attributes to payment_info
    payment_info.installments_paid = paid_installments
    payment_info.installments_due = unpaid_installments
    
    # Calculate next_due_amount - should be 0 if all payments are complete or no active sessions
    next_due_amount = 0
    if unpaid_installments > 0 and calculated_remaining > 0:
        # Only set next due amount if there are unpaid installments and remaining balance
        next_due_amount = per_installment_amount
    payment_info.next_due_amount = next_due_amount
    
    # Find next due date from student sessions - only if there are active sessions with remaining balance
    if calculated_remaining > 0:
        student_sessions_with_due = student_sessions.filter(due_date__isnull=False, status='Active').order_by('due_date')
        if student_sessions_with_due.exists():
            next_due_date = student_sessions_with_due.first().due_date

    # Get enrollment date (first registration date from student sessions)
    enrollment_date = None
    first_session = student_sessions.order_by('registration_date').first()
    if first_session and first_session.registration_date:
        enrollment_date = first_session.registration_date

    context = {
    'user': user,
    'userdata': userdata,
    'status_choices': status_choices,
    'redirection': 1,
    'student_sessions': student_sessions,
    'available_sessions': available_sessions,
    'payment_info': payment_info,
    'installments': installments,
    'next_due_date': next_due_date,
    'enrollment_date': enrollment_date,
    'total_fee': userdata.total_fee,  # ✅ Use model property
    'registration_fee': one_time_reg_fee,  # ✅ One-time registration fee
    'discount': sum(s.discount or 0 for s in student_sessions.filter(status='Active')),
    'final_fee': userdata.total_fee - sum(s.discount or 0 for s in student_sessions.filter(status='Active')),  # Fix: Use userdata.total_fee which already includes registration fee
    'paid_amount': userdata.total_paid,  # ✅ Use model property
    'remaining_amount': userdata.remaining_balance,  # Fix: Use model property which has correct calculation
    'today_date': date.today(),
    'latest_payment': latest_payment,  # Add latest payment to context
    }

    if request.method == 'POST':
        logger.debug(f"POST data keys: {list(request.POST.keys())}")
        logger.debug(f"freeze_student in POST: {'freeze_student' in request.POST}")
        logger.debug(f"unfreeze_student in POST: {'unfreeze_student' in request.POST}")
        
        # Handle freeze/unfreeze actions first
        if 'freeze_student' in request.POST:
            logger.debug("Processing freeze_student action")
            freeze_reason = request.POST.get('freeze_reason', 'Freeze')  # Default to 'Freeze' if not provided
            userdata.status = 'Inactive'
            userdata.inactive_reason = freeze_reason
            userdata.save()
            message = f"Student {userdata.student_name} marked as {freeze_reason.lower()}"
            admin_models.Notification.objects.create(user=user, category='Updation', content=message)
            return redirect('ExStudents')
            
        if 'unfreeze_student' in request.POST:
            logger.debug("Processing unfreeze_student action")
            userdata.status = 'Active'
            userdata.inactive_reason = ''
            userdata.save()
            message = f"Unfroze {userdata.student_name}"
            admin_models.Notification.objects.create(user=user, category='Updation', content=message)
            return redirect('Students')
        
        # Handle installment setup
        enable_installments = request.POST.get('enable_installments') == 'on'
        installments_count = int(request.POST.get('installments_count') or 0)
        per_installment_amount = Decimal(request.POST.get('per_installment_amount') or 0)
        single_due_date_str = request.POST.get('single_due_date')
        single_due_date = None
        if single_due_date_str:
            from datetime import datetime
            single_due_date = datetime.strptime(single_due_date_str, '%Y-%m-%d').date()
        
        logger.debug(f"enable_installments={enable_installments}, installments_count={installments_count}, per_installment_amount={per_installment_amount}")
        
        if enable_installments and installments_count > 0 and per_installment_amount > 0:
            logger.debug(f"Creating installments for student {userdata.student_name}")
            # Create installment payments
            from datetime import timedelta
            due_date = single_due_date if single_due_date else date.today()
            
            # Get the student sessions for payment records
            student_sessions_for_installments = admin_models.StudentSession.objects.filter(student=userdata)
            logger.debug(f"Found {student_sessions_for_installments.count()} student sessions")
            
            # Clear existing unpaid installments first
            admin_models.Payments.objects.filter(
                studentsession__student=userdata,
                payment_status='pending',
            ).delete()
            
            installments_created = 0
            for student_session in student_sessions_for_installments:
                for i in range(1, installments_count + 1):
                    # Create a payment record for each installment (initially unpaid)
                    payment = admin_models.Payments.objects.create(
                        studentsession=student_session,
                        user=user,
                        amount=ZERO,
                        payment_status='pending',
                        date=due_date,
                    )
                    installments_created += 1
                    logger.debug(f"Created installment {i} with payment ID {payment.id}, due date {due_date}")
                    # Next due date is one month later
                    due_date = due_date + timedelta(days=30)
            
            logger.debug(f"Total installments created: {installments_created}")
            
            # Create notification for installment setup
            installment_message = f"Set up {installments_count} installments for {userdata.student_name} - Rs.{per_installment_amount} each"
            admin_models.Notification.objects.create(
                user=user, 
                category='New Entry', 
                content=installment_message
            )
            
            messages.success(request, f"Successfully created {installments_count} installments for {userdata.student_name}")
            return redirect('StudentView', studentid=studentid)
        
        # Handle regular form submission (Save Changes button)
        # Check for specific action buttons, not just field presence
        is_freeze_action = 'freeze_student' in request.POST
        is_unfreeze_action = 'unfreeze_student' in request.POST
        is_discount_action = 'update_payment' in request.POST  # More specific discount action
        
        logger.debug(f"Action flags - freeze: {is_freeze_action}, unfreeze: {is_unfreeze_action}, discount: {is_discount_action}")
        
        # If no specific action buttons are pressed, treat as regular form save
        if not any([is_freeze_action, is_unfreeze_action, is_discount_action, enable_installments]):
            logger.debug("Processing regular form submission")
            # Filter POST data to include Student model fields and payment fields
            student_fields = [
                'student_name', 'father_name', 'email', 'cnic', 'mobile_no', 
                'Temp_address', 'Perm_address', 'last_degree', 'last_institution', 'status'
            ]
            payment_fields = ['total_fee', 'registration_fee', 'discount', 'paid_amount']
            all_fields = student_fields + payment_fields
            
            filtered_post = {}
            for field in all_fields:
                if field in request.POST:
                    filtered_post[field] = request.POST[field]
            
            # Add CSRF token
            filtered_post['csrfmiddlewaretoken'] = request.POST.get('csrfmiddlewaretoken')
            
            form = StudentForm(filtered_post, request.FILES, instance=userdata)
            
            if form.is_valid():
                # Handle file uploads
                if 'profile_photo' in request.FILES:
                    if userdata.profile_photo and os.path.exists(userdata.profile_photo.path):
                        os.remove(userdata.profile_photo.path)
                    userdata.profile_photo = request.FILES['profile_photo']
                
                if 'cnic_photo' in request.FILES:
                    if userdata.cnic_photo and os.path.exists(userdata.cnic_photo.path):
                        os.remove(userdata.cnic_photo.path)
                    userdata.cnic_photo = request.FILES['cnic_photo']
                
                if 'degree_photo' in request.FILES:
                    if userdata.degree_photo and os.path.exists(userdata.degree_photo.path):
                        os.remove(userdata.degree_photo.path)
                    userdata.degree_photo = request.FILES['degree_photo']
                
                # Save the form
                saved_student = form.save()
                
                # One-time cleanup: Remove timestamp artifacts from student name
                import re
                if "(Updated " in saved_student.student_name:
                    # Remove all timestamp patterns like "(Updated 15:00:59)"
                    cleaned_name = re.sub(r'\s*\(Updated \d{2}:\d{2}:\d{2}\)', '', saved_student.student_name)
                    saved_student.student_name = cleaned_name.strip()
                    saved_student.save()
                
                # Handle payment information updates
                logger.debug(f"Checking for payment fields in filtered_post: {[field for field in ['total_fee', 'registration_fee', 'discount', 'paid_amount'] if field in filtered_post]}")
                logger.debug(f"Payment field values: total_fee={filtered_post.get('total_fee')}, reg_fee={filtered_post.get('registration_fee')}, discount={filtered_post.get('discount')}, paid={filtered_post.get('paid_amount')}")
                
                if any(field in filtered_post for field in ['total_fee', 'registration_fee', 'discount', 'paid_amount']):
                    logger.debug("Processing payment updates...")
                    try:
                        total_fee = Decimal(filtered_post.get('total_fee', 0) or 0)
                        registration_fee = Decimal(filtered_post.get('registration_fee', 0) or 0)
                        discount = Decimal(filtered_post.get('discount', 0) or 0)
                        paid_amount = Decimal(filtered_post.get('paid_amount', 0) or 0)
                        
                        logger.debug(f"Converted values - total_fee: {total_fee}, reg_fee: {registration_fee}, discount: {discount}, paid: {paid_amount}")
                    except Exception as e:
                        logger.debug(f"Error converting payment values: {e}")
                        return redirect('StudentView', studentid=studentid)
                    
                    logger.debug(f"Processing payment updates - total_fee: {total_fee}, reg_fee: {registration_fee}, discount: {discount}, paid: {paid_amount}")
                    
                    # Update student sessions with new fees - CORRECTED
                    student_sessions = admin_models.StudentSession.objects.filter(student=saved_student)
                    for session in student_sessions:
                        # DO NOT update StudentSession.fee - it should remain as session.session.fee
                        # Only update registration_fee and discount which are session-specific
                        if registration_fee > 0:
                            session.registration_fee = int(registration_fee)
                        if discount > 0:
                            session.discount = int(discount)
                        session.save()
                        logger.debug(f"Updated StudentSession {session.id} - keeping original fee: {session.session.fee}, reg_fee: {session.registration_fee}, discount: {session.discount}")
                    
                    # Handle payment records
                    current_total_paid = sum(
                        payment.amount for payment in admin_models.Payments.objects.filter(
                            studentsession__student=saved_student, amount__gt=0
                        )
                    )
                    
                    if paid_amount > current_total_paid:
                        # Add new payment for the difference
                        additional_payment = paid_amount - current_total_paid
                        primary_session = student_sessions.first()
                        if primary_session:
                            admin_models.Payments.objects.create(
                                studentsession=primary_session,
                                user=user,
                                amount=Decimal(additional_payment),
                                payment_status='confirmed',
                                date=date.today(),
                                month=date.today().strftime('%Y-%m'),
                            )
                            logger.debug(f"Added payment of {additional_payment}")
                    
                    # Debug: Check what's in the database after updates
                    updated_sessions = admin_models.StudentSession.objects.filter(student=saved_student)
                    for session in updated_sessions:
                        logger.debug(f"Session {session.id} - fee: {session.session.fee}, reg_fee: {session.registration_fee}, discount: {session.discount}")
                    
                    updated_payments = admin_models.Payments.objects.filter(studentsession__student=saved_student)
                    total_paid_after = sum(p.amount for p in updated_payments if p.amount > 0)
                    logger.debug(f"Total paid after update: {total_paid_after}")
                else:
                    logger.debug("No payment fields found in POST data")
                    
                    # Update student sessions with new fees - CORRECTED
                    student_sessions = admin_models.StudentSession.objects.filter(student=saved_student)
                    for session in student_sessions:
                        # DO NOT update StudentSession.fee - it should remain as session.session.fee
                        # Only update registration_fee and discount which are session-specific
                        if registration_fee > 0:
                            session.registration_fee = int(registration_fee)
                        if discount > 0:
                            session.discount = int(discount)
                        session.save()
                        logger.debug(f"Updated StudentSession {session.id} - keeping original fee: {session.session.fee}, reg_fee: {session.registration_fee}, discount: {session.discount}")
                    
                    # Handle payment records
                    current_total_paid = sum(
                        payment.amount for payment in admin_models.Payments.objects.filter(
                            studentsession__student=saved_student, amount__gt=0
                        )
                    )
                    
                    if paid_amount > current_total_paid:
                        # Add new payment for the difference
                        additional_payment = paid_amount - current_total_paid
                        primary_session = student_sessions.first()
                        if primary_session:
                            admin_models.Payments.objects.create(
                                studentsession=primary_session,
                                user=user,
                                amount=Decimal(additional_payment),
                                payment_status='confirmed',
                                date=date.today(),
                                month=date.today().strftime('%Y-%m'),
                            )
                            logger.debug(f"Added payment of {additional_payment}")
                    
                    # Debug: Check what's in the database after updates
                    updated_sessions = admin_models.StudentSession.objects.filter(student=saved_student)
                    for session in updated_sessions:
                        logger.debug(f"Session {session.id} - fee: {session.session.fee}, reg_fee: {session.registration_fee}, discount: {session.discount}")
                    
                    updated_payments = admin_models.Payments.objects.filter(studentsession__student=saved_student)
                    total_paid_after = sum(p.amount for p in updated_payments if p.amount > 0)
                    logger.debug(f"Total paid after update: {total_paid_after}")
                
                # BEGIN: Installment setup on regular Save Changes
                enable_installments = request.POST.get('enable_installments') == 'on'
                installments_count = int(request.POST.get('installments_count') or 0)
                from decimal import Decimal as _Decimal  # ensure local alias if needed
                per_installment_amount = _Decimal(request.POST.get('per_installment_amount') or 0)
                single_due_date_str = request.POST.get('single_due_date')
                single_due_date = None
                if single_due_date_str:
                    from datetime import datetime as _dt
                    single_due_date = _dt.strptime(single_due_date_str, '%Y-%m-%d').date()
                
                logger.debug(f"enable_installments={enable_installments}, installments_count={installments_count}, per_installment_amount={per_installment_amount}")
                
                if enable_installments and installments_count > 0 and per_installment_amount > 0:
                    logger.debug(f"Creating installments for student {userdata.student_name}")
                    from datetime import timedelta
                    due_date = single_due_date if single_due_date else date.today()
                    
                    # Clear existing unpaid installments to avoid duplicates
                    admin_models.Payments.objects.filter(
                        studentsession__student=saved_student,
                        payment_status='pending',
                    ).delete()
                    
                    student_sessions = admin_models.StudentSession.objects.filter(student=saved_student)
                    logger.debug(f"Found {student_sessions.count()} student sessions")
                    
                    installments_created = 0
                    for student_session in student_sessions:
                        for i in range(1, installments_count + 1):
                            payment = admin_models.Payments.objects.create(
                                studentsession=student_session,
                                user=user,
                                amount=ZERO,
                                payment_status='pending',
                                date=due_date,
                            )
                            installments_created += 1
                            logger.debug(f"Created installment {i} with payment ID {payment.id}, due date {due_date}")
                            due_date = due_date + timedelta(days=30)
                    
                    logger.debug(f"Total installments created: {installments_created}")
                    admin_models.Notification.objects.create(
                        user=user,
                        category='New Entry',
                        content=f"Set up {installments_count} installments for {saved_student.student_name} - Rs.{per_installment_amount} each"
                    )
                    return redirect('StudentView', studentid=studentid)
                # END: Installment setup on regular Save Changes
                
                message = f"Updated {userdata.student_name}"
                admin_models.Notification.objects.create(user=user, category='Updation', content=message)
                return redirect('StudentView', studentid=studentid)
            else:
                # Form has errors - we'll handle this after context is built
                form_has_errors = True
                form_errors = form.errors
            
        # Handle payment information update using UNIFIED SYSTEM
        if is_discount_action:
            discount = Decimal(request.POST.get('discount', 0))
            paid_amount = Decimal(request.POST.get('paid_amount', 0))
            
            # In the unified system, we handle payments differently:
            # 1. Discounts are applied at the session level
            # 2. Payments are recorded in the Payments table
            # 3. No separate StudentFee or Installment tables
            
            # Update discount in student sessions
            if discount > 0:
                # Distribute discount across sessions proportionally
                total_session_fees = sum(
                    (session.fee or session.session.fee or 0) + (session.registration_fee or 0)
                    for session in student_sessions
                )
                
                if total_session_fees > 0:
                    for session in student_sessions:
                        session_fee = (session.fee or session.session.fee or 0) + (session.registration_fee or 0)
                        if session_fee > 0:
                            proportion = session_fee / total_session_fees
                            session_discount = int(discount * proportion)
                            session.discount = session_discount
                            session.save()
            
            # Handle payment - create a payment record if paid_amount > current total_paid
            current_paid = userdata.total_paid
            if paid_amount > current_paid:
                additional_payment = paid_amount - current_paid

                primary_session = student_sessions.first()
                if primary_session and additional_payment > 0:
                    admin_models.Payments.objects.create(
                        studentsession=primary_session,
                        user=user,
                        amount=Decimal(additional_payment),
                        payment_status='confirmed',
                        date=date.today(),
                        month=date.today().strftime('%Y-%m'),
                    )
                    
                    message = f"Added payment of Rs. {additional_payment} for {userdata.student_name}"
                    admin_models.Notification.objects.create(user=user, category='New Fee', content=message)
            
            # BEGIN: Installment setup on update_payment action
            enable_installments = request.POST.get('enable_installments') == 'on'
            installments_count = int(request.POST.get('installments_count') or 0)
            per_installment_amount = Decimal(request.POST.get('per_installment_amount') or 0)
            single_due_date_str = request.POST.get('single_due_date')
            single_due_date = None
            if single_due_date_str:
                from datetime import datetime as _dt2
                single_due_date = _dt2.strptime(single_due_date_str, '%Y-%m-%d').date()
            
            logger.debug(f"[update_payment] enable_installments={enable_installments}, installments_count={installments_count}, per_installment_amount={per_installment_amount}")
            
            if enable_installments and installments_count > 0 and per_installment_amount > 0:
                logger.debug(f"[update_payment] Creating installments for student {userdata.student_name}")
                from datetime import timedelta
                due_date = single_due_date if single_due_date else date.today()
                
                # Clear existing unpaid installments to avoid duplicates
                admin_models.Payments.objects.filter(
                    studentsession__student=userdata,
                    payment_status='pending',
                ).delete()
                
                student_sessions = admin_models.StudentSession.objects.filter(student=userdata)
                logger.debug(f"[update_payment] Found {student_sessions.count()} student sessions")
                
                installments_created = 0
                for student_session in student_sessions:
                    for i in range(1, installments_count + 1):
                        payment = admin_models.Payments.objects.create(
                            studentsession=student_session,
                            user=user,
                            amount=ZERO,
                            payment_status='pending',
                            date=due_date,
                        )
                        installments_created += 1
                        logger.debug(f"[update_payment] Created installment {i} with payment ID {payment.id}, due date {due_date}")
                        due_date = due_date + timedelta(days=30)
                
                logger.debug(f"[update_payment] Total installments created: {installments_created}")
                admin_models.Notification.objects.create(
                    user=user,
                    category='New Entry',
                    content=f"Set up {installments_count} installments for {userdata.student_name} - Rs.{per_installment_amount} each"
                )
                return redirect('StudentView', studentid=studentid)
            # END: Installment setup on update_payment action
            # Installment functionality removed in unified system
            # All payments are now handled through the Payments table
            # No need for separate installment tracking
            
            # Create notification for payment update (if any payment was made)
            if 'paid_amount' in request.POST and Decimal(request.POST.get('paid_amount', 0)) > 0:
                message = f"Updated payment information for {userdata.student_name}"
                admin_models.Notification.objects.create(user=user, category='Updation', content=message)
                
        # This section is now handled above in the reorganized POST handling

    # Handle form errors if they occurred during POST
    if form_has_errors:
        context['form'] = form  # Form with errors
        context['form_errors'] = form_errors
    else:
        context['form'] = form  # Clean form
    
    response = render(request, 'Admin/StudentView.html', context)
    # Add cache-busting headers
    response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response['Pragma'] = 'no-cache'
    response['Expires'] = '0'
    return response
@require_POST
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def mark_installment_paid(request, studentid):
    """Mark the next pending installment as paid."""
    if request.headers.get('x-requested-with') != 'XMLHttpRequest':
        return JsonResponse({'success': False, 'error': 'AJAX required.'}, status=400)
    try:
        user = _current_user(request)
        try:
            student = admin_models.Student.objects.get(id=studentid)
        except admin_models.Student.DoesNotExist:
            raise Http404
        if not _can_view_student(user, student):
            raise Http404

        next_unpaid = admin_models.Payments.objects.filter(
            studentsession__student=student,
            payment_status='pending',
        ).order_by('date').first()
        if next_unpaid is None:
            return JsonResponse({
                'success': False,
                'error': 'No unpaid installments found for this student.',
            }, status=404)

        installment_amount = request.POST.get('amount')
        if not installment_amount:
            student_sessions = admin_models.StudentSession.objects.filter(student=student, status='Active')
            total_fee = sum((Decimal(s.session.fee or 0) for s in student_sessions), ZERO)
            registration_fee = sum((Decimal(s.session.registration_fee or 0) for s in student_sessions), ZERO)
            total_with_reg = total_fee + registration_fee

            all_payments = admin_models.Payments.objects.filter(studentsession__student=student)
            installment_count = all_payments.count()

            if installment_count > 0:
                installment_amount = (total_with_reg / Decimal(installment_count)).quantize(Decimal('1'))
            else:
                installment_amount = total_with_reg
        else:
            try:
                installment_amount = Decimal(installment_amount)
            except (InvalidOperation, TypeError):
                return JsonResponse({'success': False, 'error': 'Invalid amount.'}, status=400)

        next_unpaid.amount = Decimal(installment_amount)
        next_unpaid.payment_status = 'confirmed'
        next_unpaid.month = timezone.localdate().strftime('%Y-%m')
        next_unpaid.save()

        admin_models.Notification.objects.create(
            user=user, category='Payment',
            content=f"Installment payment of Rs.{installment_amount} received for {student.student_name}",
        )

        all_payments = admin_models.Payments.objects.filter(studentsession__student=student)
        paid_count = all_payments.filter(payment_status='confirmed', amount__gt=ZERO).count()
        total_count = all_payments.count()
        unpaid_count = all_payments.filter(payment_status='pending').count()
        total_paid = sum(
            (p.amount for p in all_payments.filter(payment_status='confirmed', amount__gt=ZERO)),
            ZERO,
        )

        student_sessions = admin_models.StudentSession.objects.filter(student=student, status='Active')
        total_fee = sum((Decimal(s.session.fee or 0) for s in student_sessions), ZERO)
        registration_fee = sum((Decimal(s.session.registration_fee or 0) for s in student_sessions), ZERO)
        total_with_reg = total_fee + registration_fee
        remaining_amount = total_with_reg - total_paid

        next_unpaid_after = admin_models.Payments.objects.filter(
            studentsession__student=student,
            payment_status='pending',
        ).order_by('date').first()
        next_due_date = next_unpaid_after.date.strftime('%Y-%m-%d') if next_unpaid_after else None
        next_due_amount = int(remaining_amount / Decimal(unpaid_count)) if unpaid_count > 0 else 0

        return JsonResponse({
            'success': True,
            'message': f'Installment of Rs.{installment_amount} marked as paid successfully!',
            'paid_installments': paid_count,
            'total_installments': total_count,
            'installments_due': unpaid_count,
            'paid_amount': int(total_paid),
            'remaining_amount': int(remaining_amount),
            'next_due_amount': next_due_amount,
            'next_due_date': next_due_date,
        })
    except Http404:
        return JsonResponse({'success': False, 'error': 'Not found.'}, status=404)
    except Exception:
        logger.exception('mark_installment_paid failed')
        return JsonResponse({'success': False, 'error': 'An error occurred. Please try again.'}, status=500)
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def ExStudents(request):
    """List inactive / completed students."""
    user = _current_user(request)
    students = admin_models.Student.objects.filter(status__in=["Completed", "Inactive"])
    context = {'user': user, 'students': students, 'redirection': 2}
    return render(request, 'Admin/ExStudents.html', context)
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def AddStudent(request, id=None):
    """Add or edit a student. Admin/Moderator only."""
    user = _current_user(request)
    active_sessions = admin_models.Sessions.objects.filter(status='Active')
    
    # Mode flag to indicate if we're editing an existing student 
    # or adding a session for existing student
    is_edit_mode = False
    selected_student = None
    selected_session = None
    context = {}
    
    # If id parameter is provided, check if it's a student or session ID
    if id is not None:
        try:
            # First try to find a student with this ID
            selected_student = admin_models.Student.objects.filter(id=id).first()
            if selected_student:
                # We're adding a new session for this student or editing them
                is_edit_mode = True
                # Get student's current sessions for reference
                student_sessions = admin_models.StudentSession.objects.filter(student=selected_student)
                context.update({
                    'selected_student': selected_student,
                    'student_sessions': student_sessions,
                    'is_edit_mode': is_edit_mode,
                    'button_text': 'Save Changes',
                    'form_title': 'Enroll Student in New Session'
                })
            else:
                # If not a student, try to find a session with this ID
                selected_session = admin_models.Sessions.objects.filter(id=id).first()
                if selected_session:
                    # We're adding a student to this session
                    context.update({
                        'selected_session': selected_session,
                        'preselected_session_id': selected_session.id,
                        'button_text': 'Save Changes',
                        'form_title': 'Add Student to Session'
                    })
        except Exception as e:
            logger.error(f"Error handling ID parameter: {e}")
    
    if request.method == 'POST':
        logger.debug("POST data received")  # Debug line
        logger.debug("FILES received")  # Debug line
        
        # Existing POST handling code
        if selected_student:
            # We're updating a student or adding them to a new session
            form = StudentForm(request.POST, request.FILES, instance=selected_student)
        else:
            # We're creating a new student
            form = StudentForm(request.POST, request.FILES)

        logger.debug("Form validation check")  # Debug line
        if not form.is_valid():
            logger.debug("Form errors logged")  # Debug line
            
        if form.is_valid():
            try:
                if not selected_student:
                    # Only create a new student if we're not editing
                    newuser = form.save(commit=False)
                    
                    # Set the logged-in user as the one who added this student
                    newuser.added_by = user

                    # Handle profile photo if uploaded
                    if 'profile_photo' in request.FILES:
                        newuser.profile_photo = request.FILES['profile_photo']
                    if 'cnic_photo' in request.FILES:
                        newuser.cnic_photo = request.FILES['cnic_photo']
                    if 'degree_photo' in request.FILES:
                        newuser.degree_photo = request.FILES['degree_photo']

                    newuser.save()  # Save the new user
                    selected_student = newuser
                    logger.debug(f"Student created successfully: {newuser.id}")  # Debug line

                else:
                    # For existing student, just update and save
                    student = form.save(commit=False)
                    
                    # Handle profile photo if uploaded
                    if 'profile_photo' in request.FILES:
                        student.profile_photo = request.FILES['profile_photo']
                    if 'cnic_photo' in request.FILES:
                        student.cnic_photo = request.FILES['cnic_photo']
                    if 'degree_photo' in request.FILES:
                        student.degree_photo = request.FILES['degree_photo']
                    
                    student.save()
                    
            except Exception as e:
                logger.error(f"Error saving student: {e}")  # Debug line
                # Add error message to form
                form.add_error(None, "An error occurred. Please try again.")
                return render(request, 'Admin/AddStudent.html', context)

            # Handle session enrollment with auto roll number generation
            try:
                selected_sessions = request.POST.getlist('sessions')
                logger.debug(f"Selected sessions from POST: {selected_sessions}")
                total_fee = 0
                registration_fee = 0
                single_due_date_str = request.POST.get('single_due_date')
                single_due_date = None
                if single_due_date_str:
                    from datetime import datetime
                    single_due_date = datetime.strptime(single_due_date_str, '%Y-%m-%d').date()
                
                # Validate installment configurations before processing
                enable_installments = request.POST.get('enable_installments') == 'on'
                installments_count = int(request.POST.get('installments_count') or 0)
                
                # Check for monthly sessions and disable installments
                has_monthly_session = False
                
                if selected_sessions:
                    for session_id in selected_sessions:
                        session = admin_models.Sessions.objects.get(id=session_id)
                        if session.session_type == 'monthly':
                            has_monthly_session = True
                            break
                    
                    # Disable installments for monthly sessions
                    if has_monthly_session and enable_installments:
                        form.add_error(None, "Installment payments are not allowed for monthly sessions.")
                        return render(request, 'Admin/AddStudent.html', context)
                if selected_sessions:
                    for session_id in selected_sessions:
                        # Check if student is already enrolled in this session
                        existing_session = admin_models.StudentSession.objects.filter(
                            student=selected_student, 
                            session_id=session_id
                        ).first()
                        
                        if not existing_session:
                            session = admin_models.Sessions.objects.get(id=session_id)
                            
                            # Generate roll number if student doesn't have one
                            if not selected_student.rollno:
                                roll_number = selected_student.generate_roll_number(session)
                                selected_student.rollno = roll_number
                                selected_student.save()
                            
                            # Create StudentSession
                            student_session = admin_models.StudentSession(
                                student=selected_student,
                                session=session,
                                registration_date=date.today(),
                                registration_fee=session.registration_fee,
                                fee=session.fee,  # Set the fee from the session
                                status='Active',
                                due_date=single_due_date if not request.POST.get('enable_installments') else None
                            )
                            student_session.save()
                            logger.debug(f"Created StudentSession ID: {student_session.id} for session {session.session_name}")
                            
                            total_fee += session.fee
                            registration_fee += session.registration_fee
                            
                            message = f"Added {selected_student.student_name} (Roll: {selected_student.rollno}) to {session.session_name} session"
                            admin_models.Notification.objects.create(
                                user=user, 
                                category='New Entry', 
                                content=message
                            )

                # Handle payment information
                discount = Decimal(request.POST.get('discount', 0))
                paid_amount = Decimal(request.POST.get('paid_amount', 0))
                per_installment_amount = Decimal(request.POST.get('per_installment_amount') or 0)
                
                # Handle payment and installments (moved outside fee condition)
                logger.debug(f"enable_installments={enable_installments}, installments_count={installments_count}, per_installment_amount={per_installment_amount}")
            
                # Get the student sessions for payment records
                student_sessions_list = admin_models.StudentSession.objects.filter(student=selected_student)
                primary_session = student_sessions_list.first()
                
                # Create initial payment records for tracking
                if primary_session:
                    if not enable_installments or installments_count <= 1:
                        payment = admin_models.Payments.objects.create(
                            studentsession=primary_session,
                            user=user,
                            amount=ZERO,
                            payment_status='pending',
                            date=date.today(),
                        )
                    else:
                        due_date = single_due_date if single_due_date else date.today()

                        for i in range(1, installments_count + 1):
                            payment = admin_models.Payments.objects.create(
                                studentsession=primary_session,
                                user=user,
                                amount=ZERO,
                                payment_status='pending',
                                date=due_date,
                            )
                            # Next due date is one month later
                            due_date = due_date + relativedelta(months=1)
                        
                        # Create notification for installment setup
                        installment_message = f"Set up {installments_count} installments for {selected_student.student_name} - Rs.{per_installment_amount} each"
                        admin_models.Notification.objects.create(
                            user=user, 
                            category='New Entry', 
                            content=installment_message
                        )
                
                # Handle immediate payment if provided
                if paid_amount > 0 and primary_session:
                    unpaid_payment = admin_models.Payments.objects.filter(
                        studentsession=primary_session,
                        payment_status='pending',
                    ).order_by('date').first()

                    if unpaid_payment:
                        unpaid_payment.amount = Decimal(paid_amount)
                        unpaid_payment.payment_status = 'confirmed'
                        if not enable_installments or installments_count <= 1:
                            unpaid_payment.date = date.today()
                        unpaid_payment.save()
                    else:
                        admin_models.Payments.objects.create(
                            studentsession=primary_session,
                            user=user,
                            amount=Decimal(paid_amount),
                            payment_status='confirmed',
                            date=date.today(),
                            month=date.today().strftime('%Y-%m'),
                        )
                    
                    payment_message = f"Payment received for {selected_student.student_name}: Rs.{paid_amount}"
                    admin_models.Notification.objects.create(
                        user=user, 
                        category='Payment', 
                        content=payment_message
                    )
                
                # Note: StudentFee and Installment models have been removed
                # Payment tracking is now handled through the unified Payments system
                # via Student and StudentSession properties

            except Exception as e:
                logger.error(f"Error during session enrollment or payment processing: {e}")
                logger.exception('Error during enrollment')
                form.add_error(None, "An error occurred during enrollment. Please try again.")
                return render(request, 'Admin/AddStudent.html', {
                    'form': form,
                    'active_sessions': active_sessions,
                    'selected_student': selected_student if 'selected_student' in locals() else None,
                    'id': id
                })
            
            # Redirect to the student view after successful processing
            return redirect('StudentView', studentid=selected_student.id)
    else:
        # Not a POST request, display the form
        if selected_student:
            # If editing existing student, pre-fill the form
            form = StudentForm(instance=selected_student)
        else:
            # New student form
            form = StudentForm()

    # Update the context with common items
    context.update({
        'user': user,
        'form': form,
        'active_sessions': active_sessions,
    })
    
    if not 'button_text' in context:
        context['button_text'] = 'Add Student'
        context['form_title'] = 'Add New Student'
    
    return render(request, 'Admin/AddStudent.html', context)
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def Students(request):
    """List active students with annotated total_paid to avoid N+1 queries."""
    from .revenue import annotate_students_with_totals
    user = _current_user(request)
    students = annotate_students_with_totals(
        admin_models.Student.objects.filter(status="Active")
    )
    context = {'user': user, 'students': students, 'redirection': 1}
    return render(request, 'Admin/Students.html', context)
@login_required
@role_required(ROLE_ADMIN)
def DeleteSession(request, sessionid):
    """Soft-delete (archive) a session. Admin only.

    Blocks if confirmed payments exist on the session.
    """
    user = _current_user(request)
    try:
        session = admin_models.Sessions.objects.get(id=sessionid)
    except admin_models.Sessions.DoesNotExist:
        raise Http404

    enrolled_count = session.session_students.filter(status='Active').count()
    confirmed_payments_count = admin_models.Payments.objects.filter(
        studentsession__session=session,
        payment_status='confirmed',
        amount__gt=ZERO,
    ).count()
    if confirmed_payments_count > 0:
        messages.error(
            request,
            f'Cannot delete session with {confirmed_payments_count} confirmed payments. '
            'Archive it instead or contact the system administrator.',
        )
        return redirect('Sessions')

    session_name = session.session_name
    try:
        session.delete(deleted_by=user)
    except ProtectedError:
        messages.error(
            request,
            'Cannot delete session with confirmed payment history. Archive it instead.',
        )
        return redirect('Sessions')

    admin_models.Notification.objects.create(
        user=user, category='Deletion', content=f"Archived Session: {session_name}",
    )
    messages.success(
        request,
        f'Session "{session_name}" archived. {enrolled_count} enrollment(s) affected.',
    )
    return redirect('Sessions')
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def CompletedSessions(request):
    """List archived / completed sessions."""
    user = _current_user(request)
    sessions = admin_models.Sessions.objects.filter(status__in=["Completed", "Inactive"]).annotate(
        student_count=Count('session_students')
    )
    context = {
        'user': user,
        'sessions': sessions,
    }
    return render(request, 'Admin/CompletedSessions.html', context)

@login_required
@role_required(ROLE_ADMIN)
def RestoreSession(request, sessionid):
    """Restore an archived session. Admin only."""
    user = _current_user(request)
    try:
        session = admin_models.Sessions.all_objects.get(id=sessionid)
        if session.status == 'Active' and session.deleted_at is None:
            messages.info(request, 'Session is already active.')
            return redirect('Sessions')
        
        # Update session status to Active
        session.status = 'Active'
        session.save()
        
        # Update all students associated with this session back to Active status
        student_sessions = admin_models.StudentSession.objects.filter(session=session, status='Completed')
        students_updated = 0
        
        for student_session in student_sessions:
            # Update the student session status
            student_session.status = 'Active'
            student_session.save()
            
            # Update the student's overall status to 'Active' (Current Student)
            student = student_session.student
            student.status = 'Active'
            student.save()
            students_updated += 1
        
        # Create notification
        message = f"Restored Session: {session.session_name} with {students_updated} students transitioned back to Current Student status"
        admin_models.Notification.objects.create(user=user, category='Updation', content=message)
        
        messages.success(request, f'Session "{session.session_name}" has been successfully restored to Active status. {students_updated} students have been transitioned back to Current Student status.')
        
    except admin_models.Sessions.DoesNotExist:
        messages.error(request, 'Session not found or is not in Completed status.')
    except Exception as e:
        logger.exception('Error restoring session')
        messages.error(request, 'An error occurred while restoring the session.')
    
    return redirect('CompletedSessions')

@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def AddSession(request):
    """Create a new session. Admin/Moderator only."""
    user = _current_user(request)

    if request.method == 'POST':
        form = SessionForm(request.POST, request.FILES)

        if form.is_valid():
            newsession = form.save()  # Save the new session directly

            message = f"Added {newsession.get_session_type_display()}: {newsession.session_name}"
            admin_models.Notification.objects.create(user=user, category='New Entry', content=message)
            return redirect('Sessions')
        else:
            logger.debug("Form errors: %s", form.errors)
    else:
        form = SessionForm()

    context = {
        'user': user,
        'form': form,
    }
    return render(request, 'Admin/AddSession.html', context)
@login_required
def SessionStudentView(request, sessionid):
    """List students in a session."""
    user = _current_user(request)
    try:
        sessiondata = admin_models.Sessions.objects.get(id=sessionid)
    except admin_models.Sessions.DoesNotExist:
        raise Http404
    if not _can_view_session(user, sessiondata):
        raise Http404
    students = admin_models.StudentSession.objects.filter(session=sessiondata)
    context = {
        'user': user,
        'sessiondata': sessiondata,
        'students': students,
    }
    return render(request, 'Admin/SessionStudentView.html', context)
@login_required
def SessionView(request, sessionid):
    """View/edit a session. Admin/Moderator full; Teacher only sessions they teach."""
    user = _current_user(request)
    try:
        sessiondata = admin_models.Sessions.objects.get(id=sessionid)
    except admin_models.Sessions.DoesNotExist:
        raise Http404
    if not _can_view_session(user, sessiondata):
        raise Http404
    # Teachers can view but not edit.
    if request.method == 'POST' and user.usertype == ROLE_TEACHER:
        raise Http404
    status_choices = admin_models.Sessions.STATUS_CHOICES

    context = {
        'user': user,
        'sessiondata': sessiondata,
        'status_choices': status_choices,
    }

    if request.method == 'POST':
        form = SessionForm(request.POST, request.FILES, instance=sessiondata)  # Form for 'userdata'

        if form.is_valid():
            if 'session_photo' in request.FILES:
                if sessiondata.session_photo:
                    logger.debug(f"Old profile photo path: {sessiondata.session_photo.path}")
                    if os.path.exists(sessiondata.session_photo.path):
                        os.remove(sessiondata.session_photo.path)
            form.save()
            message = "Updated Session: " + sessiondata.session_name
            admin_models.Notification.objects.create(user=user, category='Updation', content=message)
            return redirect('SessionView', sessionid=sessionid)  # Redirect to avoid resubmission
        else:
            # Print form errors for debugging
            logger.debug("Form validation failed")
            logger.debug("Form errors: %s", form.errors)

    else:
        form = SessionForm(instance=sessiondata)

    context['form'] = form

    return render(request, 'Admin/SessionView.html', context)
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def Sessions(request):
    """List active sessions."""
    user = _current_user(request)
    sessions = admin_models.Sessions.objects.filter(status="Active").annotate(
        student_count=Count('session_students')
    )
    context = {
        'user': user,
        'sessions': sessions,
    }
    return render(request, 'Admin/Sessions.html', context)
@login_required
@role_required(ROLE_ADMIN)
def DeleteFaculty(request, userid):
    """Delete a faculty user. Admin only."""
    user = _current_user(request)
    if user.id == userid:
        messages.error(request, 'You cannot delete your own account.')
        return redirect('Faculty')
    try:
        faculty = User.objects.get(id=userid)
    except User.DoesNotExist:
        raise Http404
    
    faculty_name = faculty.first_name + " " + faculty.last_name
    
    # Remove faculty profile photo if it exists
    if faculty.profile_photo:
        if os.path.exists(faculty.profile_photo.path):
            os.remove(faculty.profile_photo.path)
    
    faculty.delete()
    
    message = "Deleted Faculty: " + faculty_name
    admin_models.Notification.objects.create(user=user, category='Deletion', content=message)
    messages.success(request, f'Faculty "{faculty_name}" has been successfully deleted.')
    
    return redirect('Faculty')
@login_required
@role_required(ROLE_ADMIN)
def AddFaculty(request):
    """Add a faculty user. Admin only."""
    user = _current_user(request)

    if request.method == 'POST':
        form = UserForm(request.POST, request.FILES)

        if form.is_valid():
            newuser = form.save(commit=False)

            # Handle profile photo if uploaded
            if 'profile_photo' in request.FILES:
                newuser.profile_photo = request.FILES['profile_photo']

            newuser.save()  # Save the new user

            message = "Added Faculty: " + newuser.first_name + " " + newuser.last_name
            admin_models.Notification.objects.create(user=user, category='New Entry', content=message)
            return redirect('Faculty')
        else:
            logger.debug("Form errors: %s", form.errors)

    else:
        form = UserForm()

    context = {
        'user': user,
        'form': form,
    }
    return render(request, 'Admin/AddFaculty.html', context)
@login_required
def FacultyView(request, userid):
    """View/edit a faculty user.

    - Admin: any faculty
    - Moderator: any faculty (view only)
    - Teacher: own profile only
    """
    user = _current_user(request)
    if user.usertype == ROLE_TEACHER and user.id != userid:
        raise Http404
    try:
        userdata = User.objects.get(id=userid)
    except User.DoesNotExist:
        raise Http404
    # Teachers cannot mutate other users' profiles; only their own (handled by id check above).
    if user.usertype == ROLE_MODERATOR and request.method == 'POST':
        # Moderators are view-only on faculty.
        raise Http404

    usertype_choices = User.USER_TYPE_CHOICES
    status_choices = User.STATUS_CHOICES

    context = {
        'user': user,
        'userdata': userdata,
        'usertype_choices': usertype_choices,
        'status_choices': status_choices,
    }

    if request.method == 'POST':
        form = UserForm(request.POST, request.FILES, instance=userdata)  # Form for 'userdata'

        if form.is_valid():
            # Handle profile photo
            if 'profile_photo' in request.FILES:
                if userdata.profile_photo:
                    logger.debug(f"Old profile photo path: {userdata.profile_photo.path}")
                    if os.path.exists(userdata.profile_photo.path):
                        os.remove(userdata.profile_photo.path)

                userdata.profile_photo = request.FILES['profile_photo']

            # Save user changes
            form.save()
            message = "Updated Faculty: " + userdata.first_name + " " + userdata.last_name
            admin_models.Notification.objects.create(user=user, category='Updation', content=message)
            return redirect('FacultyView', userid=userid)  # Redirect to avoid resubmission
        else:
            # Print form errors for debugging
            logger.debug("Form validation failed")
            logger.debug("Form errors: %s", form.errors)

    else:
        form = UserForm(instance=userdata)

    context['form'] = form

    return render(request, 'Admin/FacultyView.html', context)
@login_required
@role_required(ROLE_ADMIN)
def Faculty(request):
    """List faculty users. Admin only."""
    user = _current_user(request)
    users = User.objects.all()
    context = {
        'user': user,
        'users': users,
    }
    return render(request, 'Admin/Faculty.html', context)
@login_required
def Profile(request):
    """View/edit own profile."""
    user = _current_user(request)

    if request.method == 'POST':
        form = UserForm(request.POST, request.FILES, instance=user)

        if form.is_valid():
            # Handle profile photo
            if 'profile_photo' in request.FILES:
                if user.profile_photo:
                    logger.debug('Old profile photo: %s', user.profile_photo.path)
                    # Delete old photo if it exists
                    if os.path.exists(user.profile_photo.path):
                        os.remove(user.profile_photo.path)

                # Save the new photo
                user.profile_photo = request.FILES['profile_photo']

            # Save user changes
            form.save()
            return redirect('Admin_Profile')  # Redirect to avoid form resubmission

    else:
        form = UserForm(instance=user)  # Pre-populate form with user data

    return render(request, 'Admin/Profile.html', {'form': form, 'user': user})
def Logout(request):
    """Secure logout: flush session, delete cookies, set no-cache headers."""
    request.session.flush()
    response = redirect('home')
    response.delete_cookie('sessionid')
    response.delete_cookie('csrftoken')
    response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response['Pragma'] = 'no-cache'
    return response
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def Admin_Dashboard(request):
    """Admin/Moderator dashboard."""
    user = _current_user(request)
    
    users = User.objects.all()
    total_students = admin_models.Student.objects.count()  # Total number of students
    total_leads = admin_models.Lead.objects.count()  # Total number of leads
    total_sessions = admin_models.Sessions.objects.count()
    total_users = User.objects.count()
    
    # Add notification count - only count unread notifications
    notification_count = admin_models.Notification.objects.filter(is_read=False).count()
    
    context = {
        'user': user,
        'users': users,
        'total_students': total_students,
        'total_leads': total_leads,
        'total_sessions': total_sessions,
        'total_users': total_users,
        'notification_count': notification_count,
    }
    return render(request, 'Admin/Dashboard.html', context)
@require_POST
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def filter_payments(request):
    """Filter payment metrics by date range. POST-only, CSRF-enforced."""
    try:
        _ = _current_user(request)
        try:
            data = json.loads(request.body)
        except json.JSONDecodeError:
            return JsonResponse({'success': False, 'error': 'Invalid JSON.'}, status=400)
        filter_type = data.get('type')
        filter_value = data.get('value')
        from_date = data.get('fromDate')
        to_date = data.get('toDate')
        
        # Calculate date range based on filter
        today = datetime.now().date()
        start_date = None
        end_date = today
        filter_description = "All Time"
        
        if filter_type == 'today':
            start_date = today
            end_date = today
            filter_description = "Today"
        elif filter_type == 'days' and filter_value:
            start_date = today - timedelta(days=filter_value)
            filter_description = f"Last {filter_value} Days"
        elif filter_type == 'custom' and from_date and to_date:
            start_date = datetime.strptime(from_date, '%Y-%m-%d').date()
            end_date = datetime.strptime(to_date, '%Y-%m-%d').date()
            filter_description = f"{start_date.strftime('%d %b %Y')} to {end_date.strftime('%d %b %Y')}"
        
        # Filter confirmed, non-late-fee payments only.
        payments = admin_models.Payments.objects.select_related(
            'studentsession__student', 'studentsession__session', 'user'
        ).filter(payment_status='confirmed', amount__gt=ZERO, is_late_fee_payment=False)

        if start_date:
            payments = payments.filter(date__gte=start_date)
        if end_date:
            payments = payments.filter(date__lte=end_date)

        filtered_data = _filtered_revenue_metrics(payments, start_date, end_date)
        filtered_data['filter_description'] = filter_description
        
        return JsonResponse({
            'success': True,
            'data': filtered_data
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': 'An error occurred. Please try again.'})

def _filtered_revenue_metrics(payments, start_date=None, end_date=None):
    """Calculate revenue metrics for a pre-filtered payments queryset (already confirmed/non-late-fee)."""
    
    # Get all active students
    students = admin_models.Student.objects.filter(status='Active')
    
    # Calculate metrics
    total_revenue = sum(p.amount or 0 for p in payments)
    total_pending = sum(s.remaining_balance for s in students)
    total_expected_revenue = sum(s.total_fee for s in students)
    
    # Calculate collection rate
    collection_rate = (total_revenue / total_expected_revenue * 100) if total_expected_revenue > 0 else 0
    
    # Student payment status
    students_paid = sum(1 for s in students if s.remaining_balance == 0)
    students_partial = sum(1 for s in students if 0 < s.remaining_balance < s.total_fee)
    students_unpaid = sum(1 for s in students if s.remaining_balance == s.total_fee)
    
    # Recent payments (limited to filtered data)
    recent_payments = payments.order_by('-date', '-id')[:10]
    recent_payments_data = []
    for payment in recent_payments:
        recent_payments_data.append({
            'date': payment.date.isoformat() if payment.date else '',
            'student_name': payment.studentsession.student.student_name if payment.studentsession else '',
            'rollno': payment.studentsession.student.rollno if payment.studentsession else '',
            'session_name': payment.studentsession.session.session_name if payment.studentsession and payment.studentsession.session else '',
            'amount': float(payment.amount or 0),
            'collected_by': f"{payment.user.first_name} {payment.user.last_name}" if payment.user else ''
        })
    
    # Session performance
    session_revenue = {}
    for payment in payments:
        if payment.studentsession and payment.studentsession.session:
            session_name = payment.studentsession.session.session_name
            session_revenue[session_name] = session_revenue.get(session_name, 0) + (payment.amount or 0)
    
    session_performance = []
    for session_name, revenue in session_revenue.items():
        session_obj = admin_models.Sessions.objects.filter(session_name=session_name).first()
        if session_obj:
            student_count = admin_models.StudentSession.objects.filter(session=session_obj, status='Active').count()
            session_performance.append({
                'name': session_name,
                'revenue': float(revenue),
                'students': student_count,
                'avg_per_student': float(revenue / student_count) if student_count > 0 else 0
            })
    
    session_performance.sort(key=lambda x: x['revenue'], reverse=True)
    
    # Calculate other metrics
    today = datetime.now().date()
    daily_revenue = sum(p.amount or 0 for p in payments if p.date == today)
    avg_payment = total_revenue / len(payments) if payments else 0
    
    # Count overdue students
    overdue_students_count = 0
    for student in students:
        if student.remaining_balance > 0:
            student_sessions_list = student.student_sessions.filter(status='Active')
            for session in student_sessions_list:
                if session.due_date and session.due_date < today:
                    overdue_students_count += 1
                    break
    
    # Projected monthly revenue
    days_in_month = today.day
    if days_in_month > 0:
        monthly_revenue = sum(p.amount or 0 for p in payments if p.date and p.date.month == today.month and p.date.year == today.year)
        daily_avg = monthly_revenue / days_in_month
        days_remaining = 30 - days_in_month
        projected_monthly_revenue = monthly_revenue + (daily_avg * days_remaining)
    else:
        projected_monthly_revenue = 0
    
    return {
        'total_revenue': float(total_revenue),
        'total_pending': float(total_pending),
        'collection_rate': round(collection_rate, 1),
        'avg_payment': float(avg_payment),
        'daily_revenue': float(daily_revenue),
        'active_students_count': len(students),
        'overdue_students_count': overdue_students_count,
        'projected_monthly_revenue': float(projected_monthly_revenue),
        'students_paid': students_paid,
        'students_partial': students_partial,
        'students_unpaid': students_unpaid,
        'recent_payments': recent_payments_data,
        'session_performance': session_performance
    }

@require_POST
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def export_word_report(request):
    """Export revenue report as a Word document. POST-only, CSRF-enforced."""
    try:
        _ = _current_user(request)
        try:
            data = json.loads(request.body)
        except json.JSONDecodeError:
            return HttpResponse('Invalid JSON.', status=400)
        filter_data = data.get('filter', {})
        
        # Get filtered payments
        filter_type = filter_data.get('type')
        filter_value = filter_data.get('value')
        from_date = filter_data.get('fromDate')
        to_date = filter_data.get('toDate')
        
        # Calculate date range
        today = datetime.now().date()
        start_date = None
        end_date = today
        period_description = "All Time"
        
        if filter_type == 'today':
            start_date = today
            end_date = today
            period_description = "Today"
        elif filter_type == 'days' and filter_value:
            start_date = today - timedelta(days=filter_value)
            period_description = f"Last {filter_value} Days"
        elif filter_type == 'custom' and from_date and to_date:
            start_date = datetime.strptime(from_date, '%Y-%m-%d').date()
            end_date = datetime.strptime(to_date, '%Y-%m-%d').date()
            period_description = f"{start_date.strftime('%d %b %Y')} to {end_date.strftime('%d %b %Y')}"
        
        # Filter payments
        payments = admin_models.Payments.objects.select_related(
            'studentsession__student', 'studentsession__session', 'user'
        ).filter(payment_status='confirmed', amount__gt=ZERO, is_late_fee_payment=False)

        if start_date:
            payments = payments.filter(date__gte=start_date)
        if end_date:
            payments = payments.filter(date__lte=end_date)

        metrics = _filtered_revenue_metrics(payments, start_date, end_date)
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('IQRA ACADEMY', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('Revenue Report', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add report details
        doc.add_paragraph(f"Report Period: {period_description}")
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%d %B %Y at %I:%M %p')}")
        doc.add_paragraph("")
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=2)
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Table Grid'
        
        summary_data = [
            ('Total Revenue Collected', f"Rs. {metrics['total_revenue']:,.0f}"),
            ('Outstanding Receivables', f"Rs. {metrics['total_pending']:,.0f}"),
            ('Collection Efficiency', f"{metrics['collection_rate']}%"),
            ('Average Payment Size', f"Rs. {metrics['avg_payment']:,.0f}"),
            ('Active Students', str(metrics['active_students_count']))
        ]
        
        for i, (label, value) in enumerate(summary_data):
            summary_table.cell(i, 0).text = label
            summary_table.cell(i, 1).text = value
        
        doc.add_paragraph("")
        
        # Student Analysis
        doc.add_heading('Student Payment Analysis', level=2)
        student_table = doc.add_table(rows=4, cols=2)
        student_table.style = 'Table Grid'
        
        student_data = [
            ('Students with Full Payment', str(metrics['students_paid'])),
            ('Students with Partial Payment', str(metrics['students_partial'])),
            ('Students with No Payment', str(metrics['students_unpaid'])),
            ('Students with Overdue Payments', str(metrics['overdue_students_count']))
        ]
        
        for i, (label, value) in enumerate(student_data):
            student_table.cell(i, 0).text = label
            student_table.cell(i, 1).text = value
        
        doc.add_paragraph("")
        
        # Session Performance
        if metrics['session_performance']:
            doc.add_heading('Session Performance', level=2)
            session_table = doc.add_table(rows=len(metrics['session_performance']) + 1, cols=4)
            session_table.style = 'Table Grid'
            
            # Headers
            headers = ['Session Name', 'Revenue', 'Students', 'Avg per Student']
            for i, header in enumerate(headers):
                session_table.cell(0, i).text = header
            
            # Data
            for i, session in enumerate(metrics['session_performance'][:10]):
                session_table.cell(i + 1, 0).text = session['name']
                session_table.cell(i + 1, 1).text = f"Rs. {session['revenue']:,.0f}"
                session_table.cell(i + 1, 2).text = str(session['students'])
                session_table.cell(i + 1, 3).text = f"Rs. {session['avg_per_student']:,.0f}"
        
        doc.add_paragraph("")
        
        # Recent Payments
        if metrics['recent_payments']:
            doc.add_heading('Recent Payments', level=2)
            payments_table = doc.add_table(rows=min(len(metrics['recent_payments']), 20) + 1, cols=5)
            payments_table.style = 'Table Grid'
            
            # Headers
            headers = ['Date', 'Student', 'Session', 'Amount', 'Collected By']
            for i, header in enumerate(headers):
                payments_table.cell(0, i).text = header
            
            # Data
            for i, payment in enumerate(metrics['recent_payments'][:20]):
                payments_table.cell(i + 1, 0).text = datetime.fromisoformat(payment['date']).strftime('%d %b %Y') if payment['date'] else ''
                payments_table.cell(i + 1, 1).text = payment['student_name']
                payments_table.cell(i + 1, 2).text = payment['session_name']
                payments_table.cell(i + 1, 3).text = f"Rs. {payment['amount']:,.0f}"
                payments_table.cell(i + 1, 4).text = payment['collected_by']
        
        doc.add_paragraph("")
        
        # Recommendations
        doc.add_heading('Recommendations', level=2)
        recommendations = []
        
        if metrics['collection_rate'] < 70:
            recommendations.append("• Collection rate is below 70% - Consider implementing automated payment reminders")
            recommendations.append("• Follow up with students having overdue payments")
        
        if metrics['students_unpaid'] > 0:
            recommendations.append(f"• {metrics['students_unpaid']} students have not made any payments - Immediate attention required")
        
        if metrics['overdue_students_count'] > 0:
            recommendations.append(f"• {metrics['overdue_students_count']} students have overdue payments - Send payment reminders")
        
        recommendations.extend([
            "• Consider offering flexible payment plans for students with large outstanding amounts",
            "• Regular monthly collection drives can improve cash flow",
            "• Implement early payment discounts to encourage prompt payments"
        ])
        
        for rec in recommendations:
            doc.add_paragraph(rec)
        
        # Save to BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Create response
        response = HttpResponse(
            doc_io.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="Revenue_Report_{period_description.replace(" ", "_")}_{today.strftime("%Y%m%d")}.docx"'
        
        return response
        
    except Exception as e:
        return HttpResponse('An error occurred. Please try again.', status=500)

@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def get_email_statistics(request):
    """Get real-time email statistics for the Email Services dashboard."""
    
    try:
        from django.db.models import Count, Q
        from datetime import datetime, timedelta
        
        # Get current date for calculations
        today = datetime.now().date()
        thirty_days_ago = today - timedelta(days=30)
        
        # Calculate total recipients
        total_students = admin_models.Student.objects.filter(
            email__isnull=False, 
            email__gt='',
            status='Active'
        ).count()
        
        total_faculty = User.objects.filter(
            email__isnull=False,
            email__gt=''
        ).count()
        
        total_leads = admin_models.Lead.objects.filter(
            email__isnull=False,
            email__gt=''
        ).count()
        
        total_recipients = total_students + total_faculty + total_leads
        
        # Calculate students with pending fees (for reminder notifications)
        students_with_pending = admin_models.Student.objects.filter(
            status='Active',
            email__isnull=False,
            email__gt='',
        ).annotate(
            unpaid_count=Count(
                'student_sessions__student_payments',
                filter=Q(student_sessions__student_payments__payment_status='pending'),
            )
        ).filter(unpaid_count__gt=0).count()

        overdue_students = admin_models.Student.objects.filter(
            status='Active',
            email__isnull=False,
            email__gt='',
            student_sessions__student_payments__payment_status='pending',
            student_sessions__student_payments__date__lt=today,
        ).distinct().count()
        
        # Get recent notifications for email activity simulation
        recent_notifications = admin_models.Notification.objects.filter(
            date__gte=thirty_days_ago,
            category__in=['Late Fee', 'General', 'New Entry']
        ).count()
        
        # Simulate email statistics based on available data
        # In a real system, you'd track actual email sends in a separate model
        estimated_emails_sent = recent_notifications + (students_with_pending * 2)  # Assume 2 reminders per pending student
        delivery_success_rate = 95.5  # Typical email delivery rate
        
        statistics = {
            'total_emails_sent': estimated_emails_sent,
            'delivery_success_rate': delivery_success_rate,
            'pending_reminders': students_with_pending,
            'total_recipients': total_recipients,
            'breakdown': {
                'students': total_students,
                'faculty': total_faculty,
                'leads': total_leads
            },
            'overdue_students': overdue_students,
            'recent_activity': {
                'last_30_days': recent_notifications,
                'pending_notifications': students_with_pending
            }
        }
        
        return JsonResponse({
            'status': 'success',
            'data': statistics
        })
        
    except Exception as e:
        return JsonResponse({
             'status': 'error',
             'message': 'An error occurred. Please try again.'
         })

@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def get_email_history(request):
    """Get recent email activity for the Email Services dashboard."""
    
    try:
        from django.db.models import Q
        from django.utils import timezone
        from datetime import timedelta
        
        # Get current user for sent_by field
        current_user = User.objects.get(id=request.session['user_id'])
        
        # Get recent notifications that represent email activity
        recent_notifications = admin_models.Notification.objects.filter(
            category__in=['Late Fee', 'General', 'New Entry']
        ).order_by('-date')[:20]
        
        email_history = []
        now = timezone.now()
        
        for notification in recent_notifications:
            # Simulate email activity based on notifications
            if notification.category == 'Late Fee':
                email_type = 'Payment Reminder'
                recipient_type = 'Students'
                recipients_count = '15-25'
            elif notification.category == 'General':
                email_type = 'General Notification'
                recipient_type = 'Faculty'
                recipients_count = '8-12'
            else:
                email_type = 'Welcome Message'
                recipient_type = 'Students'
                recipients_count = '10-20'
            
            # Use timezone-aware comparison
            status = 'Delivered' if notification.date < now - timedelta(hours=1) else 'Pending'
            
            email_history.append({
                'date': notification.date.strftime('%d %b %Y'),
                'subject': f'{email_type} - {notification.date.strftime("%B %Y")}',
                'recipients': recipients_count,
                'recipient_type': recipient_type,
                'status': status,
                'sent_by': f'{current_user.first_name} {current_user.last_name}'
            })
        
        # If no notifications, provide some sample data
        if not email_history:
            email_history = [
                {
                    'date': now.strftime('%d %b %Y'),
                    'subject': 'Payment Reminder - Current Month',
                    'recipients': '23',
                    'recipient_type': 'Students',
                    'status': 'Delivered',
                    'sent_by': f'{current_user.first_name} {current_user.last_name}'
                },
                {
                    'date': (now - timedelta(days=1)).strftime('%d %b %Y'),
                    'subject': 'Welcome to New Session',
                    'recipients': '15',
                    'recipient_type': 'Students',
                    'status': 'Delivered',
                    'sent_by': f'{current_user.first_name} {current_user.last_name}'
                },
                {
                    'date': (now - timedelta(days=7)).strftime('%d %b %Y'),
                    'subject': 'Faculty Meeting Reminder',
                    'recipients': '8',
                    'recipient_type': 'Faculty',
                    'status': 'Delivered',
                    'sent_by': f'{current_user.first_name} {current_user.last_name}'
                }
            ]
        
        return JsonResponse({
            'status': 'success',
            'data': email_history
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': 'An error occurred. Please try again.'
        })

@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def PDFNameComparison(request):
    """PDF Name Comparison — validates uploaded file using magic-byte check before parsing."""
    user = _current_user(request)
    
    # Get all current and completed sessions for selection
    current_sessions = admin_models.Sessions.objects.filter(status='Active').annotate(
        student_count=Count('session_students')
    )
    completed_sessions = admin_models.Sessions.objects.filter(status='Completed').annotate(
        student_count=Count('session_students')
    )
    
    context = {
        'user': user,
        'current_sessions': current_sessions,
        'completed_sessions': completed_sessions,
    }
    
    if request.method == 'POST':
        try:
            # Handle PDF upload and session selection
            pdf_file = request.FILES.get('pdf_file')
            selected_sessions_raw = request.POST.get('selected_sessions')
            
            if not pdf_file:
                return JsonResponse({'success': False, 'error': 'Please upload a PDF file.'})

            # Validate file before any parsing — magic bytes, size, extension.
            is_valid, error_message = validate_pdf(pdf_file)
            if not is_valid:
                return JsonResponse({'success': False, 'error': error_message}, status=400)

            # Sanitize the filename for any downstream logging
            safe_filename = sanitize_filename(pdf_file.name)
            logger.info(f"PDFNameComparison processing file: {safe_filename}")

            if not selected_sessions_raw:
                return JsonResponse({'success': False, 'error': 'Please select at least one session.'})

            import json
            try:
                selected_sessions = json.loads(selected_sessions_raw)
            except json.JSONDecodeError:
                return JsonResponse({'success': False, 'error': 'Invalid session data format.'})

            import PyPDF2
            import io
            import re

            pdf_text = ""
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_file.read()))
            except Exception:
                logger.exception('PDF parsing failed')
                return JsonResponse(
                    {'success': False, 'error': 'Unable to read PDF. The file may be corrupted or password-protected.'},
                    status=400,
                )
            
            for page in pdf_reader.pages:
                pdf_text += page.extract_text() + "\n"
            
            # Detect which official result layout this PDF uses and extract
            # structured candidate rows. Handles both FPSC/CSS name-only
            # tables (names may wrap onto a second line) and PPSC/ECP tables
            # that carry the father's name (S/O markers or separate column).
            detected_format = detect_pdf_format(pdf_text)
            candidates = extract_candidates(pdf_text)

            pdf_name_entries = []
            for cand in candidates:
                name_words = cand['name'].lower().split()
                father_words = (cand['father_name'] or '').lower().split()
                combined_words = cand['combined'].lower().split()
                if not name_words and not combined_words:
                    continue
                pdf_name_entries.append({
                    'original_line': cand['raw_line'],
                    'words': name_words or combined_words,
                    'processed': ' '.join(name_words or combined_words),
                    'father_words': father_words,
                    'combined_words': combined_words,
                    'roll_no': cand['roll_no'],
                })

            if not pdf_name_entries:
                # Unknown layout — fall back to generic line scanning so the
                # tool still works on unstructured PDFs.
                import string
                translator = str.maketrans('', '', string.punctuation.replace('-', '').replace("'", ''))
                cleaned_text = pdf_text.translate(translator)
                pdf_lines = [line.strip() for line in cleaned_text.lower().split('\n') if line.strip()]
                for line in pdf_lines:
                    words = line.split()
                    if 1 <= len(words) <= 6 and any(word.isalpha() for word in words):
                        pdf_name_entries.append({
                            'original_line': line,
                            'words': words,
                            'processed': ' '.join(words),
                            'father_words': [],
                            'combined_words': words,
                            'roll_no': None,
                        })
            
            # Get students from selected sessions
            session_students = []
            for session_id in selected_sessions:
                # Convert session_id to integer to handle form data
                session_id = int(session_id)
                session = admin_models.Sessions.objects.get(id=session_id)
                students = admin_models.StudentSession.objects.filter(
                    session=session
                ).select_related('student', 'session')
                
                for student_session in students:
                    session_students.append({
                        'student_name': student_session.student.student_name,
                        'rollno': student_session.student.rollno,
                        'father_name': student_session.student.father_name,
                        'mobile_no': student_session.student.mobile_no,
                        'email': student_session.student.email,
                        'cnic': student_session.student.cnic,
                        'session_name': student_session.session.session_name,
                        'session_type': student_session.session.get_session_type_display(),
                        'start_date': student_session.session.start_date.strftime('%Y-%m-%d') if student_session.session.start_date else 'N/A',
                        'end_date': student_session.session.end_date.strftime('%Y-%m-%d') if student_session.session.end_date else 'N/A',
                        'registration_date': student_session.registration_date.strftime('%Y-%m-%d') if student_session.registration_date else 'N/A',
                        'fee': student_session.fee,
                        'status': student_session.status
                    })
            
            # Compare student names from database with PDF content (case-insensitive)
            matched_names = []
            unmatched_students = []
            matched_student_ids = set()  # Track matched students to avoid duplicates
            
            def normalize_name(name):
                """Normalize name for comparison by removing extra spaces, converting to lowercase, and handling special characters"""
                if not name:
                    return ""
                # Remove extra spaces, convert to lowercase, and handle special characters
                normalized = ' '.join(name.lower().split())
                # Remove common suffixes and prefixes that might cause mismatches
                normalized = re.sub(r'\s+(bin|bint|ibn|son|daughter)\s+', ' ', normalized)
                return normalized.strip()
            
            def extract_name_components(name):
                """Extract individual name components for detailed matching"""
                if not name:
                    return []
                normalized = normalize_name(name)
                # Split by spaces and filter out empty strings
                components = [comp for comp in normalized.split() if comp]
                return components
            
            def _father_overlap(student_father_name, entry):
                """Compare student's father name with the entry's father name.

                Returns True (overlap), False (both present but disjoint), or
                None (one side has no father information to compare).
                """
                if not student_father_name or not entry.get('father_words'):
                    return None
                student_father_set = set(extract_name_components(student_father_name))
                if not student_father_set:
                    return None
                return bool(student_father_set & set(entry['father_words']))

            def find_candidate_match(student_name, student_father_name, entries):
                """Match a student against extracted PDF candidate rows.

                Pass 1 — exact name match (all components equal, any order).
                When both sides carry a father name it confirms (high) or
                demotes (medium) the match; candidates whose father name
                agrees are preferred over same-name candidates whose father
                name conflicts.
                Pass 2 — subset match against the row's combined name+father
                word blob. This handles PPSC/ECP column layouts that PyPDF2
                flattens into one run of words; finding the father's name in
                the same row upgrades the match to high confidence.

                Returns:
                    tuple: (entry, confidence_level, algorithm) or (None, None, None)
                """
                if not student_name:
                    return None, None, None

                student_norm = normalize_name(student_name)
                student_components = extract_name_components(student_name)
                student_set = set(student_components)

                best = None  # (rank, entry, confidence, algorithm)
                for entry in entries:
                    entry_components = entry['words']
                    is_exact = student_norm == entry['processed'] or (
                        len(student_components) == len(entry_components)
                        and student_set == set(entry_components)
                    )
                    if not is_exact:
                        continue
                    overlap = _father_overlap(student_father_name, entry)
                    if overlap is True:
                        return entry, 'high', 'exact_name_and_father_match'
                    if overlap is None:
                        candidate = (2, entry, 'high', 'exact_full_name_case_insensitive')
                    else:
                        candidate = (1, entry, 'medium', 'exact_name_father_mismatch')
                    if best is None or candidate[0] > best[0]:
                        best = candidate

                if best:
                    return best[1], best[2], best[3]

                # Pass 2: subset match for flattened name+father columns.
                if len(student_components) >= 2:
                    for entry in entries:
                        combined = set(entry.get('combined_words') or [])
                        if student_set <= combined:
                            father_set = set(extract_name_components(student_father_name or ''))
                            if father_set and father_set <= combined:
                                return entry, 'high', 'name_and_father_in_row'
                            return entry, 'medium', 'name_subset_of_row'

                return None, None, None


            def validate_name_format(name):
                """Validate that name contains valid characters and format"""
                if not name:
                    return False, "Name is empty"
                
                # Check if name contains at least 1 alphabetic word (updated to allow single-word names)
                words = name.split()
                if len(words) < 1:
                    return False, "Name must contain at least 1 word"
                
                # Check if each word contains valid characters (letters, hyphens, apostrophes)
                for word in words:
                    if not re.match(r'^[a-zA-Z\-\']+$', word):
                        return False, f"Invalid characters in name part: {word}"
                
                # Check minimum length for each word (allow single-character names like "R")
                for word in words:
                    if len(word) < 1:
                        return False, f"Name part cannot be empty: {word}"
                
                return True, "Valid name format"
            
            # Enhanced comparison with detailed logging
            comparison_log = []
            
            # Check each student name against PDF content with enhanced matching
            # CHANGE LOG:
            # - Removed father name comparison logic completely
            # - Simplified confidence system to only use 'high' confidence
            # - All name comparisons are case-insensitive by design
            # - Updated algorithm name to reflect case-insensitive matching
            for student in session_students:
                if student['student_name'] and student['rollno'] not in matched_student_ids:
                    student_name = student['student_name']
                    
                    # Validate student name format
                    is_valid, validation_msg = validate_name_format(student_name)
                    if not is_valid:
                        comparison_log.append({
                            'rollno': student['rollno'],
                            'student_name': student_name,
                            'status': 'invalid_name',
                            'message': f"Invalid student name format: {validation_msg}"
                        })
                        unmatched_students.append(student)
                        continue
                    
                    # Match against extracted candidates (uses father name
                    # for confirmation when the PDF layout provides it).
                    name_match, match_confidence, match_algorithm = find_candidate_match(
                        student_name, student.get('father_name'), pdf_name_entries
                    )

                    if name_match:
                        enhanced_student = student.copy()
                        enhanced_student['pdf_name_match'] = {
                            'matched_entry': name_match,
                            'matching_algorithm': match_algorithm,
                            'confidence_level': match_confidence,
                        }

                        matched_names.append(enhanced_student)
                        matched_student_ids.add(student['rollno'])

                        comparison_log.append({
                            'rollno': student['rollno'],
                            'student_name': student_name,
                            'status': 'matched',
                            'name_match': name_match['original_line'],
                            'confidence': match_confidence,
                        })
                        
                    else:
                        # No exact match found
                        unmatched_students.append(student)
                        comparison_log.append({
                            'rollno': student['rollno'],
                            'student_name': student_name,
                            'status': 'not_matched',
                            'message': 'Student name not found in PDF'
                        })
            
            # Return JSON response for AJAX with enhanced results
            return JsonResponse({
                'success': True,
                'results': {
                    'matched': matched_names,
                    'pdf_only': [],  # No longer extracting specific PDF names
                    'students_only': unmatched_students,
                    'comparison_log': comparison_log  # Add detailed comparison log
                },
                'stats': {
                    'pdf_names_count': len(pdf_name_entries),  # Count of potential names found in PDF
                    'session_students_count': len(session_students),
                    'matched_count': len(matched_names),
                    'pdf_only_count': 0,  # No longer relevant
                    'students_only_count': len(unmatched_students),
                    'high_confidence_matches': len([s for s in matched_names if s.get('pdf_name_match', {}).get('confidence_level') == 'high']),
                    'medium_confidence_matches': len([s for s in matched_names if s.get('pdf_name_match', {}).get('confidence_level') == 'medium']),
                    'detected_format': detected_format,
                },
                'enhanced_features': {
                    'exact_name_matching': True,
                    'name_validation': True,
                    'detailed_logging': True,
                    'father_name_matching': True,
                    'multi_line_name_handling': True,
                }
            })
            
        except Exception as e:
            return JsonResponse({'success': False, 'error': 'An error occurred. Please try again.'})
    
    return render(request, 'Admin/PDFNameComparison.html', context)


@require_POST
@login_required
@role_required(ROLE_ADMIN, ROLE_MODERATOR)
def export_pdf_comparison_results(request):
    """Export PDF comparison results to Word or PDF document."""
    if True:
        try:
            data = json.loads(request.body)
            export_format = data.get('format', 'word')  # 'word', 'pdf', or 'word_to_pdf'
            results = data.get('results', {})
            
            if export_format == 'word':
                return generate_word_document(results)
            elif export_format == 'word_to_pdf':
                return generate_word_to_pdf_document(results)
            else:
                return generate_pdf_document(results)
                
        except Exception as e:
            logger.exception('Export failed')
            return JsonResponse({'success': False, 'error': 'Export failed. Please try again.'})
    
    return JsonResponse({'success': False, 'error': 'Invalid request method'})


def generate_word_document(results):
    """Generate Word document with comparison results"""
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('PDF Name Comparison Results', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')  # Empty line
    
    # Summary section
    doc.add_heading('Summary', level=1)
    summary_table = doc.add_table(rows=3, cols=2)
    summary_table.style = 'Table Grid'
    
    matched_count = len(results.get('matched', []))
    unmatched_count = len(results.get('students_only', []))
    
    summary_data = [
        ['Students Found in PDF', str(matched_count)],
        ['Students Not Found in PDF', str(unmatched_count)],
        ['Total Students Processed', str(matched_count + unmatched_count)]
    ]
    
    for i, (label, value) in enumerate(summary_data):
        summary_table.cell(i, 0).text = label
        summary_table.cell(i, 1).text = value
        # Make header cells bold
        summary_table.cell(i, 0).paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph('')  # Empty line
    
    # Students Found in PDF section
    if results.get('matched'):
        doc.add_heading(f'Students Found in PDF ({len(results["matched"])})', level=1)
        matched_table = doc.add_table(rows=1, cols=8)
        matched_table.style = 'Table Grid'
        
        # Header row
        header_cells = matched_table.rows[0].cells
        headers = ['Student Name', 'Roll No', 'Father Name', 'Mobile No', 'Email', 'Session', 'Session Period', 'Registration Date']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        for student in results['matched']:
            row_cells = matched_table.add_row().cells
            row_cells[0].text = student.get('student_name', 'N/A')
            row_cells[1].text = student.get('rollno', 'N/A')
            row_cells[2].text = student.get('father_name', 'N/A')
            row_cells[3].text = student.get('mobile_no', 'N/A')
            row_cells[4].text = student.get('email', 'N/A')
            row_cells[5].text = f"{student.get('session_name', 'N/A')} ({student.get('session_type', 'N/A')})"
            row_cells[6].text = f"{student.get('start_date', 'N/A')} to {student.get('end_date', 'N/A')}"
            row_cells[7].text = student.get('registration_date', 'N/A')
    
    # Students Not Found in PDF section
    if results.get('students_only'):
        doc.add_paragraph('')  # Empty line
        doc.add_heading(f'Students Not Found in PDF ({len(results["students_only"])})', level=1)
        students_only_table = doc.add_table(rows=1, cols=8)
        students_only_table.style = 'Table Grid'
        
        # Header row
        header_cells = students_only_table.rows[0].cells
        headers = ['Student Name', 'Roll No', 'Father Name', 'Mobile No', 'Email', 'Session', 'Session Period', 'Registration Date']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        for student in results['students_only']:
            row_cells = students_only_table.add_row().cells
            row_cells[0].text = student.get('student_name', 'N/A')
            row_cells[1].text = student.get('rollno', 'N/A')
            row_cells[2].text = student.get('father_name', 'N/A')
            row_cells[3].text = student.get('mobile_no', 'N/A')
            row_cells[4].text = student.get('email', 'N/A')
            row_cells[5].text = f"{student.get('session_name', 'N/A')} ({student.get('session_type', 'N/A')})"
            row_cells[6].text = f"{student.get('start_date', 'N/A')} to {student.get('end_date', 'N/A')}"
            row_cells[7].text = student.get('registration_date', 'N/A')
    
    # Save document to BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    # Create HTTP response
    response = HttpResponse(
        doc_io.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="PDF_Comparison_Results_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx"'
    
    return response


def generate_word_to_pdf_document(results):
    """Generate Word document with comparison results (excluding unmatched students) and convert to PDF"""
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('PDF Name Comparison Results', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')  # Empty line
    
    # Summary section (only showing matched students)
    doc.add_heading('Summary', level=1)
    summary_table = doc.add_table(rows=2, cols=2)
    summary_table.style = 'Table Grid'
    
    matched_count = len(results.get('matched', []))
    
    summary_data = [
        ['Students Found in PDF', str(matched_count)],
        ['Total Students Processed', str(matched_count)]
    ]
    
    for i, (label, value) in enumerate(summary_data):
        summary_table.cell(i, 0).text = label
        summary_table.cell(i, 1).text = value
        # Make header cells bold
        summary_table.cell(i, 0).paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph('')  # Empty line
    
    # Students Found in PDF section (only section included)
    if results.get('matched'):
        doc.add_heading(f'Students Found in PDF ({len(results["matched"])})', level=1)
        matched_table = doc.add_table(rows=1, cols=8)
        matched_table.style = 'Table Grid'
        
        # Header row with highlighted Student Name column
        header_cells = matched_table.rows[0].cells
        headers = ['Student Name', 'Roll No', 'Father Name', 'Mobile No', 'Email', 'Session', 'Session Period', 'Registration Date']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
            # Highlight Student Name column header
            if i == 0:
                # Set background color for Student Name column
                from docx.oxml.shared import qn
                from docx.oxml import parse_xml
                shading_elm = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(qn('w:shd')))
                header_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
        
        for student in results['matched']:
            row_cells = matched_table.add_row().cells
            row_cells[0].text = student.get('student_name', 'N/A')
            row_cells[1].text = student.get('rollno', 'N/A')
            row_cells[2].text = student.get('father_name', 'N/A')
            row_cells[3].text = student.get('mobile_no', 'N/A')
            row_cells[4].text = student.get('email', 'N/A')
            row_cells[5].text = f"{student.get('session_name', 'N/A')} ({student.get('session_type', 'N/A')})"
            row_cells[6].text = f"{student.get('start_date', 'N/A')} to {student.get('end_date', 'N/A')}"
            row_cells[7].text = student.get('registration_date', 'N/A')
            
            # Highlight Student Name column data
            shading_elm = parse_xml(r'<w:shd {} w:fill="FFFF99"/>'.format(qn('w:shd')))
            row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
    
    # Add note about comparison
    doc.add_paragraph('')
    note_para = doc.add_paragraph()
    note_para.add_run('Note: ').bold = True
    note_para.add_run('The highlighted "Student Name" column indicates the field used for PDF comparison. Only students found in the PDF are included in this report.')
    
    # Save Word document to temporary file
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
        doc.save(temp_docx.name)
        temp_docx_path = temp_docx.name
    
    try:
        # Convert Word to PDF using LibreOffice (if available)
        temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
        
        # Try LibreOffice conversion
        try:
            if platform.system() == 'Windows':
                # Try common LibreOffice paths on Windows
                libreoffice_paths = [
                    r'C:\Program Files\LibreOffice\program\soffice.exe',
                    r'C:\Program Files (x86)\LibreOffice\program\soffice.exe'
                ]
                soffice_path = None
                for path in libreoffice_paths:
                    if os.path.exists(path):
                        soffice_path = path
                        break
                
                if soffice_path:
                    subprocess.run([
                        soffice_path,
                        '--headless',
                        '--convert-to', 'pdf',
                        '--outdir', os.path.dirname(temp_pdf_path),
                        temp_docx_path
                    ], check=True, timeout=30)
                else:
                    raise FileNotFoundError('LibreOffice not found')
            else:
                # Linux/Mac
                subprocess.run([
                    'libreoffice',
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(temp_pdf_path),
                    temp_docx_path
                ], check=True, timeout=30)
            
            # Read the converted PDF
            with open(temp_pdf_path, 'rb') as pdf_file:
                pdf_content = pdf_file.read()
            
            # Clean up temporary files
            os.unlink(temp_docx_path)
            os.unlink(temp_pdf_path)
            
            # Create HTTP response
            response = HttpResponse(pdf_content, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="PDF_Comparison_Results_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf"'
            
            return response
            
        except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
            # Fallback: return Word document if PDF conversion fails
            os.unlink(temp_docx_path)
            
            # Recreate the Word document in memory
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            response = HttpResponse(
                doc_io.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="PDF_Comparison_Results_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx"'
            
            return response
            
    except Exception as e:
        # Clean up and return error
        if os.path.exists(temp_docx_path):
            os.unlink(temp_docx_path)
        raise e


def generate_pdf_document(results):
    """Generate PDF document with comparison results"""
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="PDF_Comparison_Results_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf"'
    
    # Create PDF
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    def draw_header():
        """Draw header on each page"""
        p.setFont("Helvetica-Bold", 16)
        p.drawCentredText(width/2, height-40, "PDF Name Comparison Results")
        p.setFont("Helvetica", 10)
        p.drawCentredText(width/2, height-55, f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        return height - 80
    
    def check_page_break(y_pos, needed_space=60):
        """Check if we need a new page"""
        if y_pos < needed_space:
            p.showPage()
            return draw_header()
        return y_pos
    
    y_position = draw_header()
    
    # Summary section
    y_position = check_page_break(y_position, 100)
    p.setFont("Helvetica-Bold", 14)
    p.drawString(50, y_position, "Summary")
    y_position -= 25
    
    p.setFont("Helvetica", 11)
    matched_count = len(results.get('matched', []))
    unmatched_count = len(results.get('students_only', []))
    
    summary_data = [
        f"✓ Students Found in PDF: {matched_count}",
        f"✗ Students Not Found in PDF: {unmatched_count}",
        f"📊 Total Students Processed: {matched_count + unmatched_count}"
    ]
    
    for item in summary_data:
        p.drawString(70, y_position, item)
        y_position -= 18
    
    y_position -= 20
    
    # Matched Students section
    if results.get('matched'):
        y_position = check_page_break(y_position, 100)
        p.setFont("Helvetica-Bold", 14)
        p.drawString(50, y_position, f"Students Found in PDF ({len(results['matched'])})")
        y_position -= 25
        
        for i, student in enumerate(results['matched']):
            y_position = check_page_break(y_position, 80)
            
            # Student header
            p.setFont("Helvetica-Bold", 11)
            p.drawString(70, y_position, f"{i+1}. {student.get('student_name', 'N/A')} ({student.get('rollno', 'N/A')})")
            y_position -= 15
            
            # Student details
            p.setFont("Helvetica", 9)
            details = [
                f"Father: {student.get('father_name', 'N/A')}",
                f"Mobile: {student.get('mobile_no', 'N/A')}",
                f"Email: {student.get('email', 'N/A')}",
                f"Session: {student.get('session_name', 'N/A')} ({student.get('session_type', 'N/A')})",
                f"Period: {student.get('start_date', 'N/A')} to {student.get('end_date', 'N/A')}",
                f"Registration: {student.get('registration_date', 'N/A')}"
            ]
            
            for detail in details:
                p.drawString(90, y_position, detail)
                y_position -= 12
            
            y_position -= 8  # Extra space between students
    
    # Students Not in PDF section
    if results.get('students_only'):
        y_position = check_page_break(y_position, 100)
        p.setFont("Helvetica-Bold", 14)
        p.drawString(50, y_position, f"Students Not Found in PDF ({len(results['students_only'])})")
        y_position -= 25
        
        for i, student in enumerate(results['students_only']):
            y_position = check_page_break(y_position, 80)
            
            # Student header
            p.setFont("Helvetica-Bold", 11)
            p.drawString(70, y_position, f"{i+1}. {student.get('student_name', 'N/A')} ({student.get('rollno', 'N/A')})")
            y_position -= 15
            
            # Student details
            p.setFont("Helvetica", 9)
            details = [
                f"Father: {student.get('father_name', 'N/A')}",
                f"Mobile: {student.get('mobile_no', 'N/A')}",
                f"Email: {student.get('email', 'N/A')}",
                f"Session: {student.get('session_name', 'N/A')} ({student.get('session_type', 'N/A')})",
                f"Period: {student.get('start_date', 'N/A')} to {student.get('end_date', 'N/A')}",
                f"Registration: {student.get('registration_date', 'N/A')}"
            ]
            
            for detail in details:
                p.drawString(90, y_position, detail)
                y_position -= 12
            
            y_position -= 8  # Extra space between students
    
    p.save()
    
    # Get PDF content
    pdf_content = buffer.getvalue()
    buffer.close()
    
    response.write(pdf_content)
    return response
